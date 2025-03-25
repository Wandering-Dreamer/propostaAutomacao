import docx
import pandas as pd
from docx.shared import Pt, Mm, Cm, Inches
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def change_table_cell(cell, background_color=None, font_color=None, font_size=None, bold=None, italic=None):

    if background_color:
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), background_color))
        cell._tc.get_or_add_tcPr().append(shading_elm)

    if font_color:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = docx.shared.RGBColor.from_string(font_color)

    if font_size:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = docx.shared.Pt(font_size)

    if bold is not None:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = bold

    if italic is not None:
        for p in cell.paragraphs:
            for r in p.runs:
                r.italic = italic

def change_table_row(table_row, background_color=None, font_color=None, font_size=None, bold=None, italic=None):
    for cell in table_row.cells:
        change_table_cell(cell, background_color=background_color, font_color=font_color, font_size=font_size,
                          bold=bold,
                          italic=italic)

df = pd.read_excel("SA38 Test.xlsx")
doc = docx.Document()

unique_values = df.iloc[:, 2].drop_duplicates()
print(unique_values)
out = df.reset_index().groupby([df.iloc[:, 2]])['index'].min().to_list()
print(out)

section = doc.sections[0]

section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

for i in out:

    t0 = doc.add_table(3, 1)
    t0.alignment = WD_TABLE_ALIGNMENT.CENTER
    t0.autofit = True
    t0.allow_autofit = True
    t0.style = 'Table Grid'
    t0.cell(0, 0).text = str(df.values[i, 2])
    t0.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    t0.cell(1, 0).text = "HU4A3AC HPE Tech Care Critical SVC"
    t0.cell(1, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    t0.cell(2, 0).text = "Vigência - De: " + str(df.values[1, 20]) + " até " + str(df.values[1, 21])
    t0.cell(2, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    change_table_row(t0.rows[0], background_color="#92D050", bold=True)
    change_table_row(t0.rows[1], background_color="#92D050", bold=True)
    change_table_row(t0.rows[2], background_color="#92D050", bold=True)

t = doc.add_table(df.shape[0]+1, 4)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = True
t.allow_autofit = True
t.style = 'Table Grid'
# convert the RHS to str
for j in range(df.shape[-1]):
    t.cell(0, 0).text = str("PART NUMBER")
    t.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.cell(0, 1).text = str("DESCRIÇÃO")
    t.cell(0, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.cell(0, 2).text = str("SERIAL")
    t.cell(0, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.cell(0, 3).text = str("QTD")
    t.cell(0, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    
# add the rest of the data frame
for i in range(df.shape[0]):
        t.cell(i+1,0).text = str(df.values[i,15])
        t.cell(i+1, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

for i in range(df.shape[0]):
        t.cell(i+1,1).text = str(df.values[i,16])
        t.cell(i+1, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

for i in range(df.shape[0]):
        t.cell(i+1,2).text = str(df.values[i,0])
        t.cell(i+1, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

for i in range(df.shape[0]):
        t.cell(i+1,3).text = str(df.values[i,14]) 
        t.cell(i+1, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

for i in range(df.shape[0]):
        t.cell(i,3).width = Pt(50)
        
for i in range(df.shape[0]):
        for j in range(4):
            if j == 0:
                t.cell(i,j).width = Mm(100)
            if j == 1:
                t.cell(i,j).width = Mm(150)
            if j == 2:
                t.cell(i,j).width = Mm(100)
            elif j == 3:
                t.cell(i,j).width = Mm(15)


change_table_row(t.rows[0], background_color="#92D050", bold=True)


# save the doc
doc.save('./test.docx')