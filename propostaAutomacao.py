from docx import *
from docx.shared import Pt, Inches, Mm
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import date
import locale
import ui

# Pag 1
doc = Document("Template Proposta Tecnica Comercial Servicos de Suporte_novo (003) (002).docx")

section = doc.sections[0]

section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

title = doc.add_heading('Resposta para ', level=0)
title.add_run(ui.cliente_final.get())
title.add_run(" de Hewlett Packard Enterprise \n\n\n\n")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title2 = doc.add_heading('Projeto: ', level=0)
title2.add_run(ui.num_contrato_final.get())
title2.alignment = WD_ALIGN_PARAGRAPH.LEFT
title2.add_run('\n\n\n\n\n\n\n\n')

p = doc.add_paragraph('São Paulo, ')
p.add_run(ui.str_dt)
p = doc.add_paragraph('Proposta Técnica Comercial ')
p.add_run(ui.ope_final.get())
p.add_run('\n\n\n')
doc.add_picture('image_page1.jpg', width=Pt(500))
doc.add_page_break()


# Pag 2

section = doc.sections[-1]
section2 = doc.add_section()

section2.left_margin = Inches(1.85)
section2.right_margin = Inches(0.5)
section2.top_margin = Inches(0.25)
section2.bottom_margin = Inches(1)

if ui.proposta_final.get() == 0:

    p2 = doc.add_paragraph('\n\n\nSão Paulo, ')
    p2.add_run(ui.str_dt)
    p2.add_run('\n\n\n')
    p2 = doc.add_paragraph(ui.cliente_final.get())
    p2.add_run(", ")
    p2.add_run('\n\n\n')
    vendedor = doc.add_paragraph()
    vendedor.add_run(ui.vendedor.get()).font.size = Pt(9)
    vendedor.add_run('\n').font.size = Pt(9)
    vendedor.add_run("Installed Base Specialist").font.size = Pt(9)
    vendedor.add_run('\n').font.size = Pt(9)
    vendedor.add_run(ui.var_telefone.get()).font.size = Pt(9)
    vendedor.add_run('\n').font.size = Pt(9)
    vendedor.add_run(ui.var_email.get()).font.size = Pt(9)
    vendedor.add_run('\n\n\n').font.size = Pt(9)
    vendedor.alignment = WD_ALIGN_PARAGRAPH.LEFT
    vendedor.paragraph_format.left_indent = Mm(-30.4)
    p21 = doc.add_paragraph('Prezados (as) Senhores(as):\n\n')
    p21.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p22 = doc.add_paragraph('Temos o prazer de apresentar a nossa proposta técnico-comercial referente ao Projeto: ')
    p22.add_run(ui.num_contrato_final.get())
    p22.add_run("""\nEstamos confiantes que as informações contidas nesta proposta possam atender suas necessidades, demonstrando desta maneira a potencialidade de nossa empresa nos termos de qualidade de produtos e serviços.""")
    p22.add_run('\nEsta proposta foi desenvolvida por Hewlett Packard Enterprise, que analisou todos os aspectos necessários para uma implementação bem sucedida.')
    p22.add_run('\nEstamos confiantes em demonstrar os benefícios de valor agregado da proposta e construir um relacionamento de negócio sólido e benéfico para ambas as partes.')
    p22.add_run('\nColocamo-nos à disposição para quaisquer esclarecimentos que se faça necessário.')
    p22.add_run('\n\n\n\nAtenciosamente,')
    p22.add_run('\n\n\n\n\n _____________________________\n\n\n')
    p22.add_run(ui.vendedor.get())
    p22.add_run('\n')
    p22.add_run("Installed Base Specialist")
    hpe = doc.add_paragraph('\n\n\n\n')
    hpe.add_run('Hewlett Packard Enterprise').bold = True
    hpe.add_run('\nAlameda Rio Negro, 750')
    hpe.add_run('\nBarueri, SP, 06454-000')
    hpe.add_run('\nBrazil')
    hpe.add_run('\nwww.hpe.com')
    hpe.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hpe.paragraph_format.left_indent = Mm(-30.4)
    doc.add_page_break()

elif ui.proposta_final.get() == 1:

    p2 = doc.add_paragraph('\n\n\nSão Paulo, ')
    p2.add_run(ui.str_dt)
    p2.add_run('\n\n\n')
    p2 = doc.add_paragraph(ui.cliente_final.get())
    p2.add_run(", ")
    p2.add_run('\n\n\n')
    vendedor = doc.add_paragraph()
    vendedor.add_run(ui.vendedor.get()).font.size = Pt(9)
    vendedor.add_run('\n').font.size = Pt(9)
    vendedor.add_run("Installed Base Specialist").font.size = Pt(9)
    vendedor.add_run('\n').font.size = Pt(9)
    vendedor.add_run(ui.var_telefone.get()).font.size = Pt(9)
    vendedor.add_run('\n').font.size = Pt(9)
    vendedor.add_run(ui.var_email.get()).font.size = Pt(9)
    vendedor.add_run('\n\n\n').font.size = Pt(9)
    vendedor.alignment = WD_ALIGN_PARAGRAPH.LEFT
    vendedor.paragraph_format.left_indent = Mm(-30.4)
    p21 = doc.add_paragraph('Prezados (as) Senhores(as):\n\n')
    p21.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p22 = doc.add_paragraph('A Hewlett Packard Brasil Ltda. (“HPE”) apresenta sua proposta técnica comercial referente ao Projeto de Suporte de Hardware e Software ')
    p22.add_run(ui.num_contrato_final.get())
    p22.add_run("\n\nEsta proposta foi desenvolvida pela HPE, que analisou todos os aspectos necessários para uma implementação e prestação de serviços, considerando o quanto requerido por este órgão.\n\nA HPE está confiante de que as informações contidas nesta proposta possam atender suas necessidades, demonstrando desta maneira a potencialidade da  empresa nos termos de qualidade de produtos e serviços.\n\nA HPE está confiante em demonstrar os benefícios de valor agregado da proposta e construir um relacionamento de negócio sólido e benéfico para ambas as partes, observadas as disposições contidas na Lei n ° 14.133/2021 no que tange a licitações e contratos da Administração Pública.\n\nColocamo-nos à disposição para quaisquer esclarecimentos que se faça necessário.")
    p22.add_run('\n\n\n\nAtenciosamente,')
    p22.add_run('\n\n\n\n\n _____________________________\n\n\n')
    p22.add_run(ui.vendedor.get())
    p22.add_run('\n')
    p22.add_run("Installed Base Specialist")
    hpe = doc.add_paragraph('\n\n\n\n')
    hpe.add_run('Hewlett Packard Enterprise').bold = True
    hpe.add_run('\nAlameda Rio Negro, 750')
    hpe.add_run('\nBarueri, SP, 06454-000')
    hpe.add_run('\nBrazil')
    hpe.add_run('\nwww.hpe.com')
    hpe.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hpe.paragraph_format.left_indent = Mm(-30.4)
    doc.add_page_break()

# Pag 3

title3 = doc.add_heading(level=0)
title3.add_run("\nAviso de Confidencialidade\n")
p3 = doc.add_paragraph()
p3.add_run("As informações contidas em todas as páginas deste documento / proposta é confidencial da Hewlett Packard Enterprise e Hewlett Packard Enterprise Company (a seguir coletivamente \"Hewlett Packard Enterprise\") e seguem para fins de avaliação. Ao receber o documento, o destinatário concorda em manter tais informações em sigilo e não reproduzir ou divulgar a qualquer pessoa fora do grupo diretamente responsável pela avaliação do conteúdo, a menos que a  Hewlett Packard Enterprise tenha autorizado. Não há obrigação de manter a confidencialidade de qualquer parte da informação que o destinatário tenha tido conhecimento sem restrições antes do recebimento deste documento, como é provado através de registos escritos, de negócios ou informações de conhecimento público sem que o destinatário tenha incorrido em faltas, ou que tenha sido recebido pelo destinatário através de uma terceira parte sem restrições.").font.size = Pt(9)
p3.add_run("\n\nEste documento contém informações sobre produtos, vendas e programas de serviço da  Hewlett Packard Enterprise que podem ser melhorados ou descontinuados a critério exclusivo da  Hewlett Packard Enterprise. A  Hewlett Packard Enterprise tem feito todos os esforços para incluir materiais aqui considerados confiáveis e relevantes para fins de avaliação de seu destinatário. Nem a Hewlett Packard Enterprise nem seus representantes dão qualquer garantia quanto à exatidão ou completude das informações. Portanto, este documento é apenas para fins informativos devendo ser considerado para os negócios da  Hewlett Packard Enterprise. Nem a  Hewlett Packard Enterprise nem seus representantes serão responsáveis sobre qualquer ato do destinatário ou de seus representantes, como resultado do uso das informações aqui fornecidas. A assinatura de um acordo definitivo ou assinatura de aceitação da proposta, por representantes autorizados das partes, será o único meio pelo qual a  Hewlett Packard Enterprise ou suas afiliadas serão vinculadas à proposta/ contrato.").font.size = Pt(9)
title31 = doc.add_heading(level = 0).add_run("Restrições de cópias entregues da Proposta\n")
p31 = doc.add_paragraph().add_run("""\nA proposta da Hewlett Packard Enterprise foi enviada em formato eletrônico no formato de arquivo PDF. Se o conteúdo dos arquivos originais forem diferentes da versão em PDF, somente o conteúdo da versão PDF será respeitado pela Hewlett Packard Enterprise.""").font.size = Pt(9)
title32 = doc.add_heading(level = 0).add_run("Esclarecimentos\n")
p32 = doc.add_paragraph().add_run("\nDúvidas ou  esclarecimentos sobre esta Política de Privacidade, entre em contato com seu representante de vendas.").font.size = Pt(9)
p33 = doc.add_paragraph().add_run("\n\n\n© Copyright 2025 Hewlett-Packard Development Company, L.P.").font.size = Pt(9)
p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
doc.add_page_break()

# Pag 4

title4 = doc.add_heading(level=0)
title4.add_run("Índice\n\n")
p4 = doc.add_paragraph()
p4.add_run('1. Resumo Executivo	                                                                                  5\n')
p4.add_run('2. Serviços de Suporte	                                                                      7\n')
p4.add_run('3. Especificações dos Níveis de Serviços	                                             11\n')
p4.add_run('4. Condições Comerciais         	                                                        21\n')
p4.add_run('5. Condições Gerais                	                                                        28\n')
p4.add_run('6. Termo de Aceite da Proposta / Pedido de Compra                               30\n')
p4.add_run('7. Anexos                	                                                                               32\n')
doc.add_page_break()

# Pag 5
title5 = doc.add_heading(level=0)
title5.add_run("1. Resumo Executivo\n")
p5 = doc.add_paragraph()
p5.add_run('A Hewlett Packard Enterprise acredita que as seguintes vantagens competitivas nos diferenciam da concorrência:\n\n')
p5.add_run('Amplo portfólio de soluções completas. ').bold = True
p5.add_run('Combinamos nossos recursos de infraestrutura, software e serviços para oferecer o que acreditamos ser o maior e mais completo portfólio de soluções empresariais do setor de TI. Nossa capacidade de oferecer uma ampla variedade de produtos de alta qualidade e serviços de suporte e consultoria de alto valor em um único pacote é um dos nossos principais diferenciais.')
p5.add_run('\nRoteiro de inovação para vários anos. ').bold = True
p5.add_run('Atuamos no setor de tecnologia e inovação há mais de 75 anos. Nosso amplo portfólio de propriedade intelectual e nossos recursos de pesquisa e desenvolvimento global fazem parte um roteiro de inovação mais amplo para ajudar organizações de todos os tamanhos em sua jornada de plataformas de tecnologia tradicional rumo aos sistemas de TI do futuro — denominados Novo Estilo de TI — que acreditamos serão caracterizados pelo aumento e pela proeminência inter-relacionada de computação em nuvem, Big Data, segurança empresarial, aplicativos e mobilidade.')
p5.add_run('\nDistribuição global e ecossistema de parceiros. ').bold = True
p5.add_run('Somos especialistas no fornecimento de soluções tecnológicas inovadoras para nossos clientes em ambientes complexos com a participação de vários países, vários fornecedores e/ou vários idiomas. Temos um dos conjuntos de recursos go-to-market mais completos do setor, incluindo um amplo ecossistema de parceiros de canal que nos permite comercializar e entregar ofertas de produtos a clientes localizados em, praticamente, qualquer lugar do mundo.')
p5.add_run('\nSoluções financeiras personalizadas. ').bold = True
p5.add_run('Desenvolvemos soluções financeiras inovadoras para facilitar a entrega de produtos e serviços a nossos clientes. Oferecemos soluções de investimento flexíveis e especialidade que ajudam os clientes e outros parceiros a criar implantações de tecnologia exclusivas com base em necessidades comerciais específicas.')
p5.add_run('\nEquipe de liderança experiente com histórico comprovado de desempenho de sucesso. ').bold = True
p5.add_run('Nossa equipe de gerenciamento apresenta um histórico comprovado de desempenho e execução. Nossa equipe de gerenciamento sênior soma mais de 100 anos de experiência na área e possui amplo conhecimento e experiência no setor de TI comercial e nos mercados em que competimos. Além disso, possuímos um amplo banco de talentos em gerenciamento e tecnologia que — acreditamos — nos oferece pipeline sem precedentes para futuros líderes e inovadores.')
p5.add_run('\nUm parceiro de transformação com a visão e a abrangência para ajudar os clientes a alcançar ótimos resultados comerciais').font.size = Pt(10)
p5.paragraph_format.line_spacing = Pt(13)
doc.add_picture('image_page5.png', width=Pt(350))
doc.add_page_break()

# Pag 6 

p6 = doc.add_paragraph()
p6.add_run('\nAo escolher a Hewlett Packard Enterprise como sua parceira comercial, o cliente ganhará um consultor experiente, comprometido e confiável. A Hewlett Packard Enterprise entende os muitos desafios associados à implantação da infraestrutura complexa de tecnologia da informação no ambiente de negócios em rápida mudança e altamente competitivo dos dias de hoje. Nossa estratégia concentra-se em ampliar nossos pontos fortes para oferecer maior valor para os clientes, seja com suporte pós-venda, um conjunto maior de ferramentas de virtualização e gerenciamento para empresas ou foco elevado em serviços. Oferecemos nosso incrível portfólio da maneira que for mais adequada para você — ajudando-o a transformar desafios em oportunidades. Nossas soluções completam essa estrutura ao combinar nossas tecnologias a seus objetivos comerciais avançados de forma holística e revolucionária. Com nossa visão, estratégia, experiência e liderança, a Hewlett Packard Enterprise se destaca claramente como parceira preferencial de negócios do cliente para o futuro. A equipe de profissionais de vendas e serviços da Hewlett Packard Enterprise tem a experiência necessária para traduzir as metas de negócios do cliente em soluções de TI que aprimoram a competitividade, geram um rápido retorno do investimento e fornecem proteção de longo prazo aos ativos por meio de um caminho de crescimento assegurado.')
p6.add_run('\n\nEm resumo, o que diferencia a Hewlett Packard Enterprise da concorrência é nossa proposta de valor para nossos clientes. A Hewlett Packard Enterprise ocupa uma posição única — sólida e invejável — com nossa combinação de ativos, adotando uma postura baseada em padrões abertos e cética em relação a plataformas, e com nossa capacidade de fornecer insights úteis para os clientes de maneira que nenhuma outra empresa pode fazer. Um conjunto integrado e global de produtos e serviços oferece muito mais — mais responsabilidade e agilidade comercial, flexibilidade, suporte econômico para as necessidades dinâmicas da TI, um retorno mais alto dos investimentos em TI e a melhor experiência total do cliente.')
p6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
doc.add_page_break()

# Pag 7, Pag 8, Pag 9, Pag 10, Pag 11

title7 = doc.add_heading(level=0)
title7.add_run("\n2. Serviços de Suporte\n")
if ui.contrato_final.get() == 0:
    title71 = doc.add_heading(level=0)
    title71.add_run('HPE Pointnext Tech Care')
    p7 = doc.add_paragraph()
    p7.add_run("Aplicavel aos contratos que contenham algum dos Números de produto abaixo: \n\nHU4A3AC; HU4A4AC; HU4A5AC; HU4A6AC; HU4A7AC; HU4A8AC; HU4A9AC; HU4B0AC; HU4B1AC; HU4B2AC; HU4B3AC; HU4B4AC; HU4B5AC; HU4B6AC; HU4B7AC;")
    p7.add_run("\n\nVISÃO GERAL DO SERVIÇO\n\n").bold = True
    p7.add_run('\nO HPE Pointnext Tech Care é a experiência de suporte operacional para os produtos de hardware e software da marca HPE (produtos HPE). Com o HPE Pointnext Tech Care, as equipes de TI podem manter-se focadas no desenvolvimento da empresa, buscando proativamente formas melhores de agir, ao invés de apenas reagir a problemas.')
    p7.add_run('\nO HPE Pointnext Tech Care vai além do suporte tradicional, permitindo o acesso direto a especialistas em produtos específicos e oferecendo orientação técnica geral para ajudar os clientes não só a reduzir riscos, mas também a buscar continuamente maneiras mais eficientes de agir. Os clientes HPE Pointnext Tech Care podem obter ajuda por meio de diversos canais, incluindo o telefone, fóruns HPE moderados com tempos de resposta definidos, registro automatizado de incidentes e um recurso de bate-papo em tempo real. O serviço permite o acesso a recursos técnicos oferecidos por profissionais experientes e com conhecimento especializado no hardware ou software dentro do contexto da carga de trabalho específica, evitando que o cliente perca tempo respondendo questões de triagem ou habilitação por vezes desnecessárias. O HPE Pointnext Tech Care vai além do suporte tradicional, oferecendo orientações técnicas gerais sobre a operação, gerenciamento e segurança do produto com suporte')
    p7.add_run('\nO HPE Support Center oferece uma experiência digital aprimorada e personalizada que ajuda os clientes a gerenciarem os seus ativos por meio do reconhecimento dos vários produtos instalados no seu ambiente e da maneira como eles interagem uns com os outros. Novas ferramentas de autoatendimento permitem que os clientes desempenhem certas atividades sem precisar abrir um incidente de suporte, e oferecem também um portal de recursos e conteúdos selecionados. O HPE Pointnext Tech Care oferece acesso a recursos HPE que ajudarão a alcançar a excelência operacional e a otimização do desempenho, desde a borda até a nuvem.')
    p7.add_run("\n\nESTRUTURA DO SERVIÇO\n\n").bold = True
    p7.add_run('O serviço HPE Pointnext Tech Care, conforme destacado a seguir, oferece um conjunto geral de recursos ao lado de recursos específicos de hardware e/ou software, com base na tecnologia suportada e no fato de o produto ser um hardware, um software ou ambos. Alguns recursos do serviço são aprimorados com o uso do HPE InfoSight1, permitindo que a Hewlett Packard Enterprise ofereça níveis cada vez mais altos de orientação técnica com o uso da telemetria fornecida. Os clientes que se registram online por meio do HPE Support Center ganham acesso a recursos digitais melhorados, permitindo uma maior comodidade de gerenciamento e envolvimento direto da HPE. Os tempos de resposta remota e no local variam com base no nível do serviço selecionado, com o maior nível de serviço oferecendo assistência adicional aos clientes em casos de interrupção.')
    p7.add_run("TABELA 1. ").bold = True
    p7.add_run("Resumo dos recursos do serviço")
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for cell in row.cells:
            cell.width = Pt(200)
    table.cell(0, 0).text = '•	Acesso pelo telefone a especialistas'
    table.cell(0, 2).text = '•	Bate-papo online com especialistas'
    table.cell(0, 4).text = '•	Respostas ao fórum dadas por especialistas'
    table.cell(0, 6).text = '•	Orientação técnica geral'
    table.cell(0, 8).text = '•	Assistência HPE InfoSight'
    table.cell(0, 1).text = '•	Alertas preventivos HPE InfoSight'
    table.cell(0, 3).text = '•	Registro de incidentes automatizado'
    table.cell(0, 5).text = '•	Biblioteca de dicas técnicas'
    table.cell(0, 7).text = '•	Acesso a informações e serviços de suporte eletrônico'
    table.cell(0, 9).text = '•	Gerenciamento de interrupções (apenas no nível de serviço de Crítico)'

elif ui.contrato_final.get() == 1:
    title73 = doc.add_heading(level=0)
    title73.add_run("HPE Pointnext Complete Care Starter Pack")
    p72 = doc.add_paragraph()
    p72.add_run('Aplicavel aos contratos que contenham algum dos Números de produto abaixo: \nH2T12AC; H2T12BC;')
    p72.add_run('\n\nDESCRIÇÃO DO SERVIÇO').bold = True
    p72.add_run('\nEsta ficha técnica da Hewlett Packard Enterprise (HPE) descreve o serviço HPE Pointnext Complete Care Starter Pack, que é um pacote de suporte que fornece serviços de gerenciamento do ambiente e relacionamento que fazem parte da oferta de serviços de suporte HPE Pointnext Complete Care. É um mecanismo para os clientes adquirirem a cobertura para cada produto do ambiente de TI juntamente ao HPE Pointnext Complete Care Add-ons. O serviço HPE Pointnext Complete Care Starter Pack, juntamente aos serviços de Add-on, oferecem uma alternativa  à aquisição de um contrato com descritivo técnico personalizado do HPE Pointnext Complete (SOW). Ao adquirirem os serviços Starter Pack, há três diferentes níveis de experiência para escolher, cada um oferecendo um conjunto diferente de entregáveis de serviços:')
    p72.add_run('\n1.	HPE Pointnext Complete Care Starter Pack Standard (padrão)')
    p72.add_run('\n2.	HPE Pointnext Complete Care Starter Pack Básico')
    p72.add_run('\n3.	HPE Pointnext Complete Care Starter Pack de Entry (Inicial)')
    p72.add_run('\nPara mais detalhes sobre os diferentes níveis de experiência, consulte a tabela de Entregáveis do Serviço HPE Pointnext Complete Care Starter Pack.')
    p72.add_run('\nOs clientes escolhem o serviço HPE Pointnext Complete Care Starter Pack mais adequado às necessidades deles e adicionam produtos ao ambiente de IT juntamente com HPE Pointnext Complete Care Add-on aplicável a cada produto. Isso expande os recursos do serviço Starter Pack, como detalhados nesta ficha técnica, a esses produtos durante o termo do serviço Starter Pack.')
    p72.add_run('\nO serviço HPE Pointnext Complete Care Starter Pack não inclui serviços técnicos proativos para produtos HPE; no entanto, os clientes podem adicionar esses serviços por meio da aquisição em separado de Créditos de Serviço HPE(HPE Support Credits). O serviço Starter Pack e os Créditos de  Serviço da HPE são ativados por meio da realização, pelo cliente, de um pedido aceitável de acordo com esta ficha técnica.')
    p72.add_run('\n\nBENEFÍCIOS DO SERVIÇO').bold = True
    p72.add_run('\nOs serviços HPE Pointnext Complete Care Starter Pack foram criados para ajudarem os clientes a atingirem consistentemente suas metas de nível de serviço e outros objetivos de negócios, oferecendo:')
    p72.add_run('\n•	Um método rápido e fácil para adquirirem serviços de suporte HPE para todo o ambiente de TI (e não por dispositivo)')
    p72.add_run('\n•	Identificação proativa de problemas e consultoria sobre mitigação de riscos')
    p72.add_run('\n•	Acesso aos especialistas da HPE, que podem amplificar os recursos do cliente com o objetivo geral de ajudar a reduzir riscos, aumentar a produtividade, lidar com pico de cargas de trabalho, projetos emergentes, assim liberando o tempo do cliente para focar nos objetivos e negócios estratégicos')
    p72.add_run('\n•	Opções flexíveis de suporte reativo')
    p72.add_run('\n•	Acesso prioritário aos especialistas da HPE que conhecem o ambiente do cliente e podem ajudar a lidar rapidamente com quaisquer problemas críticos')
    p72.add_run('\n•	Opções flexíveis de suporte proativo prestado por especialistas da HPE, que complementam os recursos do cliente e ajudam-o a se concentrar em outras prioridades')
    p72.add_run('\n•	Tecnologias e ferramentas remotas avançadas projetadas para reduzirem o tempo de inatividade e aumentarem a produtividade')
    p72.add_run('\n•	Uma equipe de conta designada, focada no ambiente de TI e nos objetivos de negócios do cliente e que fornece um único ponto de contato na HPE ajudando a garantir que o relacionamento do cliente com a HPE atenda às expectativas do cliente e verificando o fornecimento de todas as opções de serviços, conforme acordado')

elif ui.contrato_final.get() == 2:
    title72 = doc.add_heading(level=0)
    title72.add_run("HPE Pointnext Complete Care Add-on  ")
    p71 = doc.add_paragraph()
    p71.add_run('Aplicável aos contratos que contenham algum dos Números de produto abaixo:\nH2T05AC; H0GT7AC; HA158AC; HA156AC; HL935AC; HA162AC; HA167AC; H7G18AC; H0JC1AC; HA360AC; H2T05AC; HU4A2AC')
    p71.add_run('\n\nDESCRIÇÃO DO SERVIÇO').bold = True
    p71.add_run('\n\nEsta ficha técnica do serviço HPE Pointnext Complete Care Add-on descreve o serviço que acrescenta novos produtos a um ambiente em que o cliente esteja adquirindo um novo HPE Pointnext Complete Care ou acrescentando há um ambiente existente.')
    p71.add_run('\nPara ser elegível à aquisição do serviço HPE Pointnext Complete Care Add-on, o cliente já deve ter um ambiente com HPE Pointnext Complete Care das duas seguintes formas:')
    p71.add_run('\n•	Ter um contrato já existente com os Serviços HPE Pointnext Complete Care com um descritivo técnico (SOW) personalizado, que inclui cobertura para o ambiente em que os novos produtos serão adicionados.')
    p71.add_run('\n•	Ter adquirido um novo HPE Pointnext Complete Care Starter Pack para o ambiente ao qual os produtos serão adicionados.')
    p71.add_run('\nAs opções acima são permitidas para realização de um pedido quando:')
    p71.add_run('\n•	Para o serviço HPE Pointnext Complete Care Add-on adquirido para produtos adicionados a um contrato com um descritivo técnico(SOW) personalizado existente: Os recursos de serviço fornecidos para produtos adicionados através de HPE Pointnext Complete Care Add-on serão os estabelecidos no seu contrato personalizado existente; ou ')
    p71.add_run('\n•	 Para os produtos adicionados HPE Pointnext Complete Care Add-on adquiridos com os Serviços HPE Pointnext Complete Care Starter Pack: Os entregáveis dos serviços fornecidos para produtos adicionados através dos HPE Pointnext Complete Care Add-on serão os definidos nesta ficha técnica, conforme especificado abaixo.')
    p71.add_run('\nQuando se adquire o serviços HPE Pointnext Complete Care Add-on, estende-se a equipe de contas designada liderada por um gerente de suporte de conta (ASM, account support manager) treinado ou Consultor de Serviço(Service Advisor), gerenciamento de incidente aprimorado (EIM, enhanced incidente management) e gerenciamento de relacionamento do serviço (SRM, service relationship management). ')
    p71.add_run('\n\nAlém disso e dependendo do tipo de produto sendo adicionado ao ambiente do HPE Pointnext Complete Care do cliente, várias opções proativas específicas da tecnologia ou outras opções de suporte podem ser incluídas em uma das duas maneiras a seguir:')
    p71.add_run('\n•	Recursos de suporte proativo específicos da tecnologia ou outras opções de suporte incluídas no serviço contratual do HPE Pointnext Complete Care descritos no seu contrato personalizado (SOW) do HPE Pointnext Complete Care.')
    p71.add_run('\n•	Recursos de suporte proativo específicos da tecnologia ou outras opções de suporte adquiridas separadamente por meio dos Créditos de Serviço HPE e descritas na ficha técnica do serviço HPE Pointnext Complete Care.')
    p71.add_run('\n\nNo serviço HPE Pointnext Complete Care Add-on inclui opções de nível de serviços reativos para    cobrir os requisitos de suporte do cliente, desde os ambientes mais básicos até os mais críticos para os negócios.')
    p71.add_run('\nAo adquirir um serviço de  HPE Pointnext Complete Care Add-on juntamente ao serviço HPE Pointnext Complete Care Starter Pack, consulte a ficha técnica do serviço HPE Pointnext Complete Care para conhecer as opções de cobertura disponíveis para esses serviços de HPE Pointnext Complete Care Add-on,   incluindo os níveis de serviços reativos aplicáveis e quaisquer limitações e exclusões, conforme estabelecidas e incorporadas a esta ficha técnica e aplicáveis à sua aquisição do serviço de HPE Pointnext Complete Care Add-on.')
    p71.add_run('\n\nBENEFÍCIOS DO SERVIÇO').bold = True
    p71.add_run('\n\nO serviço HPE Pointnext Complete Care Add-on foi criado para ajudar os clientes a atingirem consistentemente seus objetivos de negócios oferecendo:')
    p71.add_run('\n•	Uma solução de suporte modular com eficiência de custos e feita sob medida para os requisitos e o ambiente exatos do cliente')
    p71.add_run('\n•	Uma equipe de conta designada, focada no ambiente de TI e nos objetivos de negócios do cliente, fornecendo um único ponto de contato na HPE, ajudando a garantir que o relacionamento do cliente com a HPE atenda às expectativas dele e verificando o fornecimento de todas as opções de serviços, conforme acordado')
    p71.add_run('\n•	Identificação proativa de problemas e consultoria sobre mitigação de riscos')
    p71.add_run('\n•	Acesso aos especialistas da HPE, que podem amplificar os recursos do cliente com o objetivo geral de ajudar a reduzir riscos, aumentar a produtividade, lidar com as cargas de trabalho de pico e projetos emergentes e liberar o tempo do cliente para objetivos de negócios estratégicos')
    p71.add_run('\n•	Opções flexíveis de suporte reativo')
    p71.add_run('\n•	Acesso prioritário aos especialistas da HPE que conhecem o ambiente do cliente e podem ajudar a lidar rapidamente com quaisquer problemas críticos')
    p71.add_run('\n•	Opções flexíveis de suporte proativo prestado por especialistas da HPE, que complementam os recursos do cliente e podem liberar os clientes para se concentrarem em outras prioridades')
    p71.add_run('\n•	Tecnologias e ferramentas remotas avançadas projetadas para reduzirem o tempo de inatividade e aumentarem a produtividade')


# pag 12, 13, 14
# if rts

if ui.rts_final.get() == 0:

    title8 = doc.add_heading(level=0)
    title8.add_run('\n\n2.4.	Vistoria de Hardware / RTS')
    p8 = doc.add_paragraph()
    p8.add_run('\n\nEste serviço é aplicável a qualquer equipamento HEWLETT PACKARD ENTERPRISE  ou equipamento suportado pela HEWLETT PACKARD ENTERPRISE  que nunca tiveram seus equipamentos cobertos por contrato de suporte de hardware, ou que já tiveram e cancelaram seu contrato e que pretendem novamente assiná-lo.')
    p8.add_run('\n\nPara incluir equipamentos que estavam sem cobertura nos últimos 45 dias é necessário que se avalie a elegibilidade do equipamento, garantido:')
    p8.add_run('\n•	Não esteja obsoleto, seja produto suportado pela HEWLETT PACKARD ENTERPRISE  ')
    p8.add_run('\n•	Esteja atualizado com as últimas configurações e revisões')
    p8.add_run('\n•	Esteja operando sem falhas conforme determinado pela HPE')
    p8.add_run('\n\nPara incluir equipamentos que estavam sem cobertura nos últimos 45 dias é necessário que se avalie a elegibilidade do equipamento, garantido:')
    p8.add_run('\n•	Pagamento da taxa RTS')
    p8.add_run('\n•	Carência de 30 dias, isto é, o equipamento do cliente não estará suportado pelos serviços HEWLETT PACKARD ENTERPRISE ora contratados. Durante o período de carência, e uma vez evidenciado a necessidade, o cliente terá direito aos serviços de suporte mediante o pagamento de um “chamado avulso / per call service”, o qual será faturado em complemento ao(s) valore(s) especificado(s) nesta proposta e em estrita observância à lista de preços praticada pela HPE.')
    p8.add_run('\nO faturamento e pagamento referente ao valor do RTS deverá ser efetuado na sua totalidade no primeiro mês de vigência do contrato de suporte.')
    p8.add_run('\n\nCollaborative Support').bold = True
    p8.add_run('\n\nAplicavel aos contratos que contenham algum dos Números de produto abaixo:\nHL935AC; HU4A1AC')
    p8.add_run('\nOs Serviços de Hardware da HEWLETT PACKARD ENTERPRISE fornecem assistência remota e se durante o atendimento, a HEWLETT PACKARD ENTERPRISE determinar que um problema é causado por software de terceiros e que não se aplica nenhuma da correções conhecidas e disponíveis conforme definido no Software Básico, o time de suporte, quando autorizado pelo Cliente, poderá chamar o fornecedor de software de terceiros seguindo os acordos de suporte contratado entre o cliente e o fornecedor.')
    p8.add_run('\nComo parte do processo de gerenciamento do chamado, a HEWLETT PACKARD ENTERPRISE fornecerá documentação e análise realizada para que o fornecedor siga com atendimento junto ao cliente.')
    p8.add_run('\nUma vez que o fornecedor do software está envolvido, o chamado aberto junto a HEWLETT PACKARD ENTERPRISE será fechado e poderá ser reaberto, se necessário a qualquer momento fazendo referência ao número original de identificação da chamada.')
    p8.add_run('\n\n2.5.	Suporte de Software').bold = True
    p8.add_run('\n\nAplicavel aos contratos que contenham algum dos Números de produto abaixo:\nHA156AC; HU1R5AC; HV2X8AC; HA158AC')
    p8.add_run('\nO Suporte de Software da HEWLETT PACKARD ENTERPRISE  fornece serviços abrangentes para produtos de software HEWLETT PACKARD ENTERPRISE  e de terceiros selecionados. Com este serviço, sua equipe de TI conta com acesso rápido e confiável às Centrais de Atendimento da HEWLETT PACKARD ENTERPRISE . Os analistas da Central de Atendimento trabalharão com sua equipe para fornecer orientações sobre as características e utilização, diagnósticos e resolução de problemas, identificação de defeitos e acesso a patches dos produtos de software.')
    p8.add_run('\n\nO Serviço também disponibiliza atualizações de software para produtos da HEWLETT PACKARD ENTERPRISE  e de terceiros elegíveis suportados pela HEWLETT PACKARD ENTERPRISE , patches de software e manuais de referência, incluindo licença de uso e cópia de novas versões de produtos de software em todos os sistemas suportados e cobertos pela licença original do mesmo.')
    p8.add_run('\nO serviço também fornece acesso eletrônico às informações de suporte, permitindo que qualquer membro de sua equipe de TI localize informações essenciais disponíveis sobre produtos e suporte. Para produtos de terceiros, este acesso está sujeito à disponibilidade  de tais informações eletrônicas por parte do fornecedor.')
    p8.add_run('\n\nPrincipais Características dos Serviços de Software').bold = True
    p8.add_run('\n\n•	Suporte remoto\n•	Acesso a recursos técnicos\n•	Análise e resolução de problemas\n•	Gerenciamento de escalação\n•	Isolamento de problemas\n•	Suporte de orientação à instalação')
    p8.add_run('\n•	Atualizações de software da HEWLETT PACKARD ENTERPRISE  e de terceiros selecionados a um custo previsível\n•	Redução dos custos de aquisição de atualizações individuais de software, devido à economia substancial de assinaturas')
    p8.add_run('\n•	Notificações automáticas sobre a disponibilidade de novas versões de software\n•	Opção de janelas de cobertura\n•	Acesso a informações e serviços eletrônicos avançados de suporte que aumentam a produtividade:')
    p8.add_run('\n–	Hewlett Packard Enterprise Support Center (https://h20564.www2.hpe.com/ ): É um site de suporte inovador onde os profissionais de TI podem obter informações sobre software e documentações, abertura eletrônica e acompanhamento de chamados, chat direto com os engenheiros de suporte da HEWLETT PACKARD ENTERPRISE , dentre outros.')
    p8.add_run('\n\n2.5.1.	Serviços de Suporte de Software Limitado (por Incidentes)').bold = True
    p8.add_run('\n\nO Suporte Técnico de Software por Incidentes fornece serviços abrangentes de suporte remoto para produtos de software selecionados de terceiros (Microsoft, Linux Red Hat ou Suse Enterprise Edition e Novell).')
    p8.add_run('\nCom este serviço, sua equipe de TI conta com acesso rápido e confiável às Centrais de Atendimento da HEWLETT PACKARD ENTERPRISE . Os analistas da Central de Atendimento trabalharão com sua equipe para fornecer orientações sobre as características e utilização, diagnósticos e resolução de problemas, identificação de defeitos e acesso a patches dos produtos de software.')
    p8.add_run('\nAlém disso, esta modalidade permite que o serviço de suporte seja adequado às necessidades de cada ambiente com opções de 10, 25, 50 ou 75 incidentes por ano, que podem ser utilizados para diversos equipamentos. ')
    p8.add_run('\n\n2.6.	Atualização de Software / RTS ').bold = True
    p8.add_run('\\n serviço é aplicável a qualquer equipamento HEWLETT PACKARD ENTERPRISE  e foi desenvolvido especificamente para clientes HEWLETT PACKARD ENTERPRISE  que nunca tiveram seus equipamentos cobertos por contrato de suporte de software, ou que já tiveram e cancelaram seu contrato e que pretendem novamente assiná-lo.')
    p8.add_run('\n\nO Serviço de Atualização de Software HEWLETT PACKARD ENTERPRISE  (RTS - Return-to-Support) é destinado a atualizar a versão de software de qualquer produto HEWLETT PACKARD ENTERPRISE  para a última versão corrente. Após o recebimento  e instalação desta nova versão, o equipamento estará qualificado para que seja feito um contrato de manutenção de software.')
    p8.add_run('\n\nA HEWLETT PACKARD ENTERPRISE  irá prover ao cliente a versão atualmente suportada, de sistema operacional, subsistemas ou softwares aplicativos HEWLETT PACKARD ENTERPRISE . A versão corrente suportada é definida como versão de software que consta na lista de preços HEWLETT PACKARD ENTERPRISE  no momento da compra do serviço RTS. O cliente deve optar para quais produtos deseja receber a atualização, e caso a versão do sistema operacional instalada no seu equipamento apareça como a mais atual, ele deve optar somente pelos subsistemas ou softwares aplicativos.')
    p8.add_run('\n\nComprando o RTS o cliente irá receber atualização dos softwares constantes no item Configuração e Preços.')
    p8.add_run('\nOs serviços de RTS deverão estar incluídos no item Configuração e Preços através dos códigos de produtos UC255AC ou UC256AC descritos como SW Updates Return to Support.')
    p8.add_run('\n\nServiços não incluídos:	\n•	Manuais de software ou atualização de manuais.\n•	Software Status Bulletins não recebidos.\n•	Auxílio na instalação das novas versões ')
    p8.add_run('\nTanto o serviço de instalação das novas versões quanto os manuais devem ser adquiridos a parte, aos preços vigentes da época da compra do Serviço HEWLETT PACKARD ENTERPRISE -RTS. ')
    p8.add_run('\nO faturamento e pagamento referente ao valor do RTS deverá ser efetuado na sua totalidade no primeiro mês de vigência do contrato de suporte.')
    doc.add_page_break()

# Pag 15


if ui.servico_final.get() == 0:
    title91 = doc.add_heading(level=0)
    title91.add_run("3.	Especificações dos Níveis de Serviços")
    title91.add_run("\nHPE Pointnext Tech Care")
    p9 = doc.add_paragraph()
    p9.add_run("Aplicável aos contratos que contenham algum dos Números de produto abaixo: \nHU4A3AC; HU4A4AC; HU4A5AC; HU4A6AC; HU4A7AC; HU4A8AC; HU4A9AC; HU4B0AC; HU4B1AC; HU4B2AC; HU4B3AC; HU4B4AC; HU4B5AC; HU4B6AC; HU4B7AC; ")
    p9.add_run("\n\nRECURSOS GERAIS").bold = True
    p9.add_run("\n\nTABELA 2.").bold = True
    p9.add_run("Resumo das opções de níveis de serviço.\n")
    table2 = doc.add_table(rows=3, cols=2)
    table2.style = 'Table Grid'
    table2.autofit = False
    table2.allow_autofit = False
    for row2 in table2.rows:
        for cell2 in row2.cells:
            cell2.width = Pt(200)
    table2.cell(0, 0).text = "(Critical) - PNs abaixo: \nHU4A3AC\nHU4A4AC\nHU4A5AC"
    table2.cell(1, 0).text = "(Essential) - PNs abaixo:\nHU4A6AC\nHU4A7AC\nHU4A8AC\nHU4A9AC\nHU4B0AC\nHU4B1AC"
    table2.cell(2, 0).text = "(Basic) - PNs abaixo: \nHU4B2AC\nHU4B3AC\nHU4B4AC\nHU4B5AC\nHU4B6AC\nHU4B7AC"
    table2.cell(0, 1).text = "Resposta em 15 minutos, 24x7, para incidentes com severidade 1 (conecte-se diretamente aos especialistas nos produtos, quando houver disponibilidade) \nGerenciamento de interrupções para incidentes de severidade 1 Compromisso de reparo de hardware em 6 horas, 24x7 (se necessário)"
    table2.cell(1, 1).text = "Resposta em 15 minutos, 24x7, para incidentes com severidade 1 (conecte-se diretamente aos especialistas nos produtos, quando houver disponibilidade) \nAtendimento no local em 4 horas, 24x7."
    table2.cell(2, 1).text = "Resposta em 2 horas, 9x5 (horário comercial padrão) Atendimento no local no próximo dia útil."
    p91 = doc.add_paragraph()
    p91.add_run("\nTodos os níveis de serviço oferecem acesso 24x7 a recursos de autoatendimento e autorresolução de problemas, registro de incidentes 24x7 e, para os dispositivos compatíveis, análises HPE InfoSight e submissão automatizada de incidentes 24x7. As opções de nível de serviço HPE Pointnext Tech Care destacadas são dependentes do produto. A HPE fornecerá os recursos de suporte a hardware para os produtos de hardware cobertos e os recursos de suporte a software para os produtos de software cobertos.")
    p91.add_run("\nAlguns recursos dos serviços talvez não estejam disponíveis em alguns idiomas e localidades. Todos os períodos de cobertura estão sujeitos à disponibilidade local. A elegibilidade do produto pode variar. Entre em contato com um escritório local de vendas da HPE ou com o representante de vendas HPE para obter informações detalhadas sobre a disponibilidade do serviço e a elegibilidade do produto.")
    p91.add_run("\n\nTABELA 3.").bold = True
    p91.add_run("Recursos gerais do serviço\n")  
    table3 = doc.add_table(rows=8, cols=2)
    table3.style = 'Table Grid'
    table3.autofit = False
    table3.allow_autofit = False
    for row3 in table3.rows:
        for cell3 in row3.cells:
            cell3.width = Pt(200)
    table3.cell(0, 0).text = "Recurso"
    table3.cell(1, 0).text = "Acesso pelo telefone a especialistas"
    table3.cell(2, 0).text = "Bate-papo online com especialistas"
    table3.cell(3, 0).text = "Respostas ao fórum dadas por especialistas"
    table3.cell(4, 0).text = "Orientação técnica geral"
    table3.cell(5, 0).text = "Assistência HPE InfoSight"
    table3.cell(6, 0).text = "Alertas preventivos do HPE InfoSight"
    table3.cell(7, 0).text = "Registro automatizado de incidentes"
    table3.cell(0, 1).text = "Especificações de fornecimento"
    table3.cell(1, 1).text = "Os clientes podem entrar em contato com o suporte HPE pelo telefone 24 horas por dia, 7 dias por semana, para registrar incidentes de suporte. O tempo de resposta dependerá do nível de serviço do produto coberto\nResposta aprimorada em 15 minutos, 24x7 (níveis de serviço Critical e Essential)\nPara incidentes de severidade 1, a HPE procura conectar o cliente a um especialista no produto ou ligar novamente para o cliente em até 15 minutos. Para outros incidentes, a HPE pode conectar o cliente a um especialista no produto ou ligar novamente para o cliente em até uma hora.\nResposta padrão em 2 horas (nível de serviço Basic)\nPara ligações sobre produtos cobertos por um contrato de serviço básico, a HPE fornecerá uma resposta por telefone em até 2 horas, realizada por um especialista durante a janela de cobertura. A disponibilidade pode variar para determinados produtos. Consulte hpe.com/services/expertchat para obter detalhes ou entre em contato com o representante de vendas HPE local.\n"
    table3.cell(2, 1).text = "Os clientes podem iniciar um bate-papo online com um recurso técnico especialista para fazer perguntas, obter ajuda ou orientação técnica geral. O bate papo online com especialistas é oferecido para que os clientes obtenham respostas rápidas sobre questões técnicas relacionadas aos seus produtos HPE. \nQuestões complexas que exijam respostas detalhadas podem ser elevadas a incidentes de suporte conforme a necessidade. O bate-papo online com especialistas é limitado apenas ao idioma inglês e está disponível durante a janela de cobertura do serviço. A disponibilidade pode variar para determinados produtos. Consulte hpe.com/services/expertchat para obter detalhes ou entre em contato com o representante de vendas HPE local.\n"
    table3.cell(3, 1).text = "Os clientes podem postar perguntas e problemas, ou discutir o uso dos produtos dentro dos fóruns da comunidade HPE. Os especialistas em produtos HPE respondem em dois dias úteis a qualquer questão não resolvida que for levantada dentro do fórum oficial da comunidade HPE para produtos cobertos pelos serviços de suporte da HPE. \nNos casos de postagens sobre tópicos que deveriam ser abordados por processos de suporte padrão, a HPE solicita que um incidente de suporte formal seja criado e segue os processos padrão de gerenciamento de incidentes da HPE. A resposta do recurso técnico especialista é limitada apenas ao idioma inglês e exige que o usuário seja registrado no HPE Support Center e esteja associado aos contratos de serviço.\n"
    table3.cell(4, 1).text = "A HPE se empenha para oferecer orientação técnica geral relativa às questões e perguntas específicas sobre os tópicos destacados a seguir acerca da operação e gerenciamento dos produtos dos clientes cobertos pelo HPE Pointnext Tech Care. A orientação técnica geral está disponível pelo telefone, internet e bate-papo, está sujeita à janela de cobertura do serviço constante no contrato de serviço e será tratada como um incidente severidade 3.\nQuando for relacionada aos assuntos detalhados ou descritos a seguir, a HPE identificará documentos, vídeos e artigos da base de conhecimento para ajudar nos tópicos abordados.\nAlém de qualquer limitação ou exclusão estabelecida nesta ficha técnica, qualquer orientação técnica geral da HPE será fornecida especificamente para os tópicos detalhados aqui e apenas para os produtos HPE cobertos por esses serviços:\n	Uso ou procedimentos corretos para usar os recursos dos produtos\nAssistência na identificação de documentação relevante ou artigos da base de conhecimento\nConselhos sobre as melhores práticas da HPE para ajudar você a gerenciar e manter os seus produtos	\nNavegação básica para usar a interface de gerenciamento do produto\nConselhos sobre as opções de gerenciamento de capacidade com base nas tendências de uso dos produtos (quando disponível)\nOrientações sobre a configuração geral do produto, que pode incluir recomendações de melhores práticas com base na experiência operacional da HPE\nOrientação sobre os passos potenciais para ajudar a trazer o produto para uma configuração compatível\nOs tópicos de orientação técnica geral mencionados anteriormente podem não ser aplicáveis a todos os produtos de hardware e/ou software cobertos por este serviço.\n"
    table3.cell(5, 1).text = "Para produtos HPE compatíveis com o HPE InfoSight (lista disponível no link a seguir), a HPE oferece suporte e orientação para a preparação, configuração e uso do HPE InfoSight.\nAdicionalmente, para esses produtos conectados, a HPE expande sua orientação técnica geral para incluir análises do HPE InfoSight, bem como os alertas e as recomendações fornecidas. Para produtos HPE configurados, mediante solicitação, a HPE oferece assistência aos clientes para que entendam os problemas, alertas e informações oferecidas pelo HPE InfoSight.\nNos casos em que as análises oferecerem recomendações incluídas nos insights da carga de trabalho HPE InfoSight, a HPE pode oferecer a qualificação da análise, recomendações e os melhores próximos passos gerais em acordo com as orientações técnicas gerais.\nPara mais informações sobre o HPE InfoSight, cobertura de dispositivos e recursos, visite infosight.hpe.com.\n"
    table3.cell(6, 1).text = "Para os produtos HPE cobertos por um contrato de serviço, conectado a e conforme compatibilidade com o HPE InfoSight: Os clientes ganham acesso a rotinas automatizadas de monitoramento que podem identificar problemas potenciais usando assinaturas regras e determinações exclusivas da HPE. \nPara problemas identificados pelo HPE InfoSight, o HPE InfoSight alerta os clientes sobre os problemas e identifica oportunidades para ações corretivas e, de acordo com a sua criticidade, pode automaticamente submeter os incidentes à HPE com as informações diagnosticadas para acelerar os diagnósticos e reparos. \nOs recursos podem variar de acordo com o produto; os dispositivos devem ser compatíveis com o HPE InfoSight, e exige-se conectividade com o HPE InfoSight.\nNos casos em que os clientes configuram o HPE InfoSight para os produtos HPE compatíveis cobertos pelo HPE Pointnext Tech Care, eles ganham acesso aos recursos analíticos aprimorados do HPE InfoSight, que oferecem insights de produtos e alertas de problemas detalhados, além de oportunidades de uso e configuração.\n"
    table3.cell(7, 1).text = "Para produtos HPE compatíveis usando as ferramentas de serviço proprietárias HPE (incluindo o HPE InfoSight), e quando estiverem conectados, os dispositivos podem submeter os incidentes à HPE automaticamente, com informações diagnósticas que podem acelerar os diagnósticos e reparos\nOnde o monitoramento e submissão automática de incidentes identificar problemas críticos que exijam o envolvimento da HPE, a HPE procura responder ao contato do cliente previamente identificado dentro da janela de cobertura do serviço, de acordo com as definições do nível de serviço adquirido. \nCaso o cliente não esteja disponível no momento do contato, ou caso solicite, a HPE agendará o acompanhamento para o próximo dia útil. Todos os problemas não-críticos serão acompanhados no próximo dia útil. Os clientes podem, a qualquer momento, de acordo com o seu nível de serviço, entrar em contato com a HPE para solicitar a continuidade do diagnóstico e da resolução do problema.\nPara mais informações, visite hpe.com/services/getconnected.\n"

    title92 = doc.add_heading(level=0)
    title92.add_run("\n3.2.	ZONAS DE DESLOCAMENTO ")
    p92 = doc.add_paragraph()
    p92.add_run("\n\nTodos os tempos de resposta do Serviço de Troca e com presença no local para hardware se aplicam somente a áreas situadas dentro de uma distância de 160 km de um centro de suporte designado pela HPE. A viagem de descolamento dentro de um raio de 320 km de distância de um centro de suporte designado pela HPE é oferecida sem custos adicionais. Se o local estiver situado a mais de 320 km do centro de suporte designado pela HPE, haverá uma cobrança adicional. As zonas de deslocamentos e os custos, se aplicáveis, podem variar em algumas regiões. Os custos de postagem (em caso de troca de peças), se aplicáveis,  podem variar em algumas regiões. Os tempos de resposta para locais situados a mais de 160 km de um centro de suporte designado pela HPE serão modificados para incluir o tempo de deslocamento, conforme mostra a tabela a seguir. ")
    p92.add_run("\n\nTABELA 9. ").bold = True
    p92.add_run("Zonas de deslocamento (exceto nível de serviço crítico)\n\n")
    table4 = doc.add_table(rows=6, cols=3)
    table4.style = 'Table Grid'
    table4.autofit = True
    table4.allow_autofit = True
    for row4 in table4.rows:
        for cell4 in row4.cells:
            cell4.width = Pt(200)
    table4.cell(0, 0).text = "Distância do ponto de suporte determinado pela HPE"
    table4.cell(0, 1).text = "Tempo de resposta para Essencial/Essencial Exchange"
    table4.cell(0, 2).text = "Tempo de resposta para Basic/Basic Exchange"
    table4.cell(1, 0).text = "0 a 80 km"
    table4.cell(1, 1).text = "4 horas"
    table4.cell(1, 2).text = "Cobertura no dia seguinte"
    table4.cell(2, 0).text = "81 a 160 km"
    table4.cell(2, 1).text = "4 horas"
    table4.cell(2, 2).text = "Cobertura no dia seguinte"
    table4.cell(3, 0).text = "161 a 320 km"
    table4.cell(3, 1).text = "8 horas"
    table4.cell(3, 2).text = "Mais 1 dia de cobertura"
    table4.cell(4, 0).text = "321 a 480 km"
    table4.cell(4, 1).text = "Estabelecido no momento do pedido e sujeito à disponibilidade"
    table4.cell(4, 2).text = "Mais 2 dias de cobertura"
    table4.cell(5, 0).text = "Mais de 480 km"
    table4.cell(5, 1).text = "Estabelecido no momento do pedido e sujeito à disponibilidade"
    table4.cell(5, 2).text = "Estabelecido no momento do pedido e sujeito à disponibilidade"
    p93 = doc.add_paragraph()
    p93.add_run("\n\nO compromisso de tempo de reparo está disponível para localidades situados até 80 km de um centro de suporte designado pela HPE. Para locais entre 81 e 160 km de tal ponto de suporte, o compromisso de tempo de reparo de hardware será ajustado de acordo com a tabela a seguir. O compromisso de tempo de reparo de hardware não está disponível para locais a mais de 160 km de um ponto de suporte designado pela HPE.")
    p93.add_run("\n\nTABELA 10. ").bold = True
    p93.add_run("Zonas de deslocamento para nível de serviço crítico\n\n")
    table5 = doc.add_table(rows=4, cols=2)
    table5.style = 'Table Grid'
    table5.autofit = False
    table5.allow_autofit = False
    for row5 in table5.rows:
        for cell5 in row5.cells:
            cell5.width = Pt(180)
    table5.cell(0, 0).text = "Distância do ponto de suporte designado pela HPE"
    table5.cell(1, 0).text = "0 a 80 km"
    table5.cell(2, 0).text = "81 a 160 km"
    table5.cell(3, 0).text = "Mais de 160 km"
    table5.cell(0, 1).text = "Tempo de reparo crítico"
    table5.cell(1, 1).text = "6 horas"
    table5.cell(2, 1).text = "8 horas"
    table5.cell(3, 1).text = "Não disponível"

elif ui.servico_final.get() == 1:

    section = doc.sections[-1]
    section3 = doc.add_section()

    section3.left_margin = Inches(0.75)
    section3.right_margin = Inches(0.75)
    section3.top_margin = Inches(1)
    section3.bottom_margin = Inches(1)
    
    title93 = doc.add_heading(level = 0)
    title93.add_run("3.	Especificações dos Níveis de Serviços")
    title93.add_run("\nHPE Pointnext Complete Care Starter Pack ")
    p94 = doc.add_paragraph()
    p94.add_run("\nAplicavel aos contratos que contenham algum dos Números de produto abaixo:\nH2T12AC; H2T12BC;")
    p94.add_run("\n\n3.3.1.	Destaques dos Entregáveis do Serviço ").bold = True
    p94.add_run("\n\nO HPE Pointnext Complete Care Starter Pack oferece três níveis de experiência de gerenciamento e relacionamento, incluindo recursos HPE designados que compreendem os objetivos de negócios de TI do cliente e trabalham para garantir que essas necessidades sejam atendidas.")
    p94.add_run("\nSujeitos a quaisquer limitações definidas neste documento, os serviços HPE Pointnext Complete Care Starter Pack Standard(padrão), básico e Entry(Inicial) estabelecem os seguintes entregáveis.")
    p94.add_run("\nTABELA 1.").bold = True
    p94.add_run("Destaques dos entregáveis do serviço HPE Pointnext Complete Care Starter Pack")
    table6 = doc.add_table(rows=15, cols=4)
    table6.style = 'Table Grid'
    table6.autofit = True
    table6.allow_autofit = True
    for row6 in table6.rows:
        for cell6 in row6.cells:
            cell6.width = Pt(180)
    table6.cell(0, 0).text = "\nEntregáveis\n"
    table6.cell(0, 1).text = "\nEspecificações de fornecimento\n"
    table6.cell(1, 0).text = "\nNível de experiência\n"
    table6.cell(1, 1).text = "\nStandard (Padrão)\n"
    table6.cell(1, 2).text = "\nBásico\n"
    table6.cell(1, 3).text = "\nEntry (Inicial)\n"
    table6.cell(2, 0).text = "\nEquipe de conta designada\n"
    table6.cell(2, 1).text = "\nA equipe de conta designada pela HPE é o representante do cliente e o ponto focal técnico e/ou operacional para o HPE Pointnext Complete Care. Esses recursos coordenam os entregáveis do HPE Pointnext Complete Care. Isso inclui serviços proativos opcionais, bem como o monitoramento de problemas, correções e consultoria que possam impactar o ambiente do cliente. Além disso, esses recursos oferecem atividades determinadas conforme detalhado mais especificamente abaixo.\n"
    a = table6.cell(2, 1)
    b = table6.cell(2, 3)
    c = table6.cell(2, 0)
    d = table6.cell(3, 0)
    a.merge(b)
    c.merge(d)
    table6.cell(3, 1).text = "\nA HPE designa os seguintes recursos de conta à organização do cliente:\n•	Gerente de suporte de conta (ASM)\n•	Gerente técnico de conta (TAM)\n•	Técnico designado ao cliente (ACE)\n"
    table6.cell(3, 2).text = "\nA HPE designa os seguintes recursos de conta à organização do cliente:\n•	ASM\n•	Consultor de serviço (Service Advisor)\n"
    table6.cell(3, 3).text = "\nA HPE designa o seguinte recurso de conta à organização do cliente:\n•	Consultor de serviço(Service Advisor)\n"
    table6.cell(4, 0).text = "\nPlanejamento de suporte à conta\n"
    table6.cell(4, 1).text = "\nEsse recurso de serviço oferece o desenvolvimento de um plano de suporte da conta (ASP) pelo ASM em parceria com a equipe de TI do cliente. \nO ASP documenta os suportes reativo e proativo adquiridos, os dispositivos, a cobertura geográfica e quaisquer outros aspectos do suporte. \nO ASP também detalha os contatos,  as funções, responsabilidades dos recursos envolvidos e fluxo de informações, que serão confirmadas com o cliente. O ASP receberá atualizações  proativamente pelo ASM conforme requirido.\n"
    table6.cell(4, 2).text = "\nIgual ao HPE Pointnext Complete Care padrão exceto pela frequência de atualização do ASP  semestralmente.\n"
    table6.cell(4, 3).text = "\nIgual ao HPE Pointnext Complete Care padrão exceto que o ASP é desenvolvido pelo Consultor do Serviço e atualizado semestralmente.\n"
    table6.cell(5, 0).text = "\nGerenciamento da implementação do serviço\n"
    table6.cell(5, 1).text = "\nNo começo do período de cobertura do suporte do serviço HPE Pointnext Complete Care, um gerente de implementação gerencia a implmentação do serviço. \nIsso pode incluir o contato com o cliente e a apresentação do cliente à equipe de conta designada, ou a equipe da conta pode agir como facilitadora enquanto o gerente de implementação coordena as atividades em segundo  plano.\n"
    table6.cell(5, 2).text = "\nMesmo que o HPE Pointnext Complete Care padrão.\n"
    table6.cell(5, 3).text = "\nMesmo que o HPE Pointnext Complete Care padrão.\n"
    table6.cell(6, 0).text = "\nGerenciamento de inventário\n"
    table6.cell(6, 1).text = "\nNo começo do período de cobertura do serviço HPE Pointnext Complete Care, a equipe de conta designada vai desenvolver e documentar um inventário com todos os produtos cobertos no ambiente do HPE Pointnext Complete Care do cliente. \nA equipe de conta da HPE designada ajudará a gerenciar as alterações com o cliente continuamente para manter atualizado esse inventário durante o período de cobertura  do serviço. Isso foi pensado para ajudar o cliente a garantir que todos os produtos para os quais o cliente deseja ter suporte no ambiente do serviço HPE Pointnext Complete Care sejam cobertos.\n"
    table6.cell(6, 2).text = "\nMesmo que o HPE Pointnext Complete Care padrão, fornecido pela equipe de conta designada.\n"
    table6.cell(6, 3).text = "\nMesmo que o HPE Pointnext Complete Care padrão, fornecido pelo Consultor do Serviço.\n"
    table6.cell(7, 0).text = "\nPlanejamento e revisão do serviço\n"
    table6.cell(7, 1).text = "\nEssas sessões trimestrais de análise fornecem um fórum aberto de comunicação para ajudar o cliente a compartilhar os objetivos de negócios e TI de sua organização. Durante essas sessões de análise, os recursos de conta da HPE podem compartilhar as melhores práticas da HPE e fornecer conselhos relacionados às necessidades e projetos operacionais atuais e futuros do cliente. \nOutros recursos HPE podem participar dessas reuniões conforme determinado pelo ASM ou Consultor do Serviço.\n"
    table6.cell(7, 2).text = "\nEssas sessões de análise semestrais fornecem um fórum aberto de comunicação para ajudar o cliente a compartilhar os objetivos de negócios e TI de sua organização. Durante essas sessões de análise, os recursos de conta da HPE podem compartilhar as melhores práticas da HPE e fornecer conselhos relacionados às necessidades e projetos operacionais atuais e futuros do cliente. \nOutros recursos HPE podem participar dessas reuniões conforme determinado pelo ASM ou Consultor do Serviço.\n"
    table6.cell(7, 3).text = "\nEssas sessões de análise semestrais fornecem um fórum aberto de comunicação para ajudar o cliente a compartilhar os objetivos de negócios e TI de sua organização. Durante essas sessões de análise, os recursos de conta da HPE podem compartilhar as melhores práticas da HPE e fornecer conselhos relacionados às necessidades e projetos operacionais atuais e futuros do cliente. \nOutros recursos HPE podem participar dessas reuniões conforme determinado pelo Consultor do Serviço.\n"
    table6.cell(8, 1).text = "\nEssas sessões de análise fornecem um fórum aberto de comunicação para ajudar o cliente a compartilhar os objetivos de negócios e TI de sua organização. Durante essas sessões de análise, os recursos de conta da HPE podem compartilhar as melhores práticas da HPE e fornecer conselhos relacionados às necessidades e projetos operacionais atuais e futuros do cliente. \nOutros recursos HPE podem participar dessas reuniões conforme determinado pelo ASM ou Consultor do Serviço.\n"
    a = table6.cell(7, 0)
    b = table6.cell(8, 0)
    c = table6.cell(8, 1)
    d = table6.cell(8, 3)
    a.merge(b)
    c.merge(d)
    table6.cell(9, 0).text = "\nAnálise das atividades de suporte\n"
    table6.cell(9, 1).text = "\nA HPE fornece ao cliente um relatório trimestral de revisão das atividades de suporte que documenta as informações de incidente de suporte reativo durante esse período específico.\nO relatório também pode destacar os fatores de risco potenciais e incluir recomendações sugeridas pela HPE.\n"
    table6.cell(9, 2).text = "\nN/D (ver Relatório de incidente)\n"
    table6.cell(9, 3).text = "\nN/D\n"
    table6.cell(10, 0).text = "\nRelatório de incidente\n"
    table6.cell(10, 1).text = "\nN/D (essas informações são incluídas no Relatório de revisão da atividade de suporte)\n"
    table6.cell(10, 2).text = "\nA HPE oferece um relatório de incidente semestral que documenta as informações de incidente de suporte reativo durante o período especificado.\n"
    table6.cell(10, 3).text = "\nN/D\n"
    table6.cell(11, 0).text = "\nConsultoria técnica e operacional\n"
    table6.cell(11, 1).text = "\nTrabalhando com o cliente, a equipe de conta designada da HPE tem uma função ativa no fornecimento de consultoria e orientação relacionadas ao fornecimento rotineiro dos serviços relacionados ao ambiente do cliente coberto pelo HPE Pointnext Complete Care. \nEsse entregável de serviço destina-se a fornecer uma breve orientação aos clientes. Solicitações substanciais de assistência, conforme determinada pela HPE a seu critério, estão fora do escopo desse serviço, mas podem ser adquiridas usando os Créditos de Serviço HPE ou Dia Técnico(Team Day). \nO ASM informará o cliente quando  uma solicitação exigir créditos ou dias técnicos e o valor necessário.\n"
    table6.cell(11, 2).text = "\nN/D (ver Assistência operacional)\n"
    table6.cell(11, 3).text = "\nN/D (ver Assistência operacional)\n"
    table6.cell(12, 0).text = "\nAssistência operacional\n"
    table6.cell(12, 1).text = "\nN/D (incluído como parte da Consultoria operacional e técnica)"
    table6.cell(12, 2).text = "\nTrabalhando com o cliente, a equipe de conta designada da HPE fornecerá assistência operacional básica ao fornecimento rotineiro dos serviços relacionados ao ambiente do cliente coberto pelo HPE Pointnext Complete Care. Esse recurso de serviço destina-se a fornecer uma breve orientação aos clientes. \nSolicitações substanciais de assistência, conforme determinadas pela HPE a seu critério, estão fora do escopo desse serviço, mas podem ser adquiridas usando os Créditos  de Serviço HPE. O Consultor do Serviço informará o cliente quando uma solicitação exigir créditos e a quantia necessária.\n"
    table6.cell(12, 3).text = "\nMesmo que o HPE Pointnext Complete Care básico.\n"
    table6.cell(13, 0).text = "\nCentro de Suporte HPE\n"
    table6.cell(13, 1).text = "\nA HPE oferece um abrangente recurso on-line para serviços, ferramentas e conhecimento disponíveis. Esse ponto único completo de TI oferece ferramentas de autodiagnóstico, assistência personalizada, ajuda e fóruns on-line, além de acesso a determinados conteúdos abrangentes de TI multifornecedores e multiplataformas.\n"
    a = table6.cell(13, 1)
    b = table6.cell(13, 3)
    a.merge(b)
    table6.cell(14, 0).text = "\nAssistência e planejamento educacionais HPE\n"
    table6.cell(14, 1).text = "\nMediante requisição do cliente, o ASM pode conduzir uma análise de alto nível das necessidades de treinamento e desenvolvimento do cliente. \nO ASM também pode fornecer assistência para entrar em contato com a equipe da HPE Education Services. O cliente pode acessar currículos de treinamento e descrições detalhadas do curso no site dos HPE Education Services na web em hpe.com/ww/learn.\n"
    table6.cell(14, 2).text = "\nMesmo que o HPE Pointnext Complete Care padrão.\n"
    table6.cell(14, 3).text = "\nMesmo que o HPE Pointnext Complete Care padrão exceto que a assistência é fornecida pelo Consultor do Serviço\n"

    p95 = doc.add_paragraph()
    p95.add_run("\n\n3.3.2.	Créditos de Serviço e Dias Técnicos ")
    p95.add_run("\nAplicavel aos contratos que contenham algum dos Números de produto abaixo (Sendo “X” a quantidade de anos”)\n• HPE Service Credit Service 10 Credits SVC: H0JD4Ax \n• HPE Service Credit Service 30 Credits SVC: H0JD5Ax")
    p95.add_run("\nSujeitos a quaisquer limitações estabelecidas neste documento, os seguintes recursos adquiríveis separadamente estão disponíveis com o serviço HPE Pointnext Complete Care Starter Pack.")
    p95.add_run("\nTABELA 2. ").bold = True
    p95.add_run("Créditos de Serviço HPE e Dia Técnico HPE\n")
    table7 = doc.add_table(rows=3, cols=2)
    table7.style = 'Table Grid'
    table7.autofit = True
    table7.allow_autofit = True
    for row7 in table7.rows:
        for cell7 in row7.cells:
            cell7.width = Pt(180)
    table7.cell(0, 0).text = "Recurso"
    table7.cell(0, 1).text = "Especificações de fornecimento"
    table7.cell(1, 0).text = "Créditos de Serviço HPE"
    table7.cell(1, 1).text = "Os clientes do serviço HPE Pointnext Complete Care Starter Pack podem adquirir separadamente os Créditos de Serviço HPE, que podem ser usados para uma variedade de serviços técnicos ajudando a manter e otimizar proativamente os produtos no ambiente do serviço HPE Pointnext Complete Care Starter Pack. \nO cliente tem a flexibilidade de escolher uma atividade do menu de Créditos de Serviço técnico pré-definido ou de trabalhar com o recurso da conta atribuído (seja ASM ou Consultor de Serviço, como aplicável) para definir uma atividade personalizada baseada nas necessidades do cliente. \nMais informações sobre limitações de serviço, responsabilidades do cliente, disposições gerais, exclusões, termos e condições podem ser encontradas na Ficha técnica dos Créditos de Serviço HPE. A ficha técnica será aplicada a qualquer compra de tais créditos e é incorporada a este documento como referência. \nOs Créditos de Serviço HPE elegíveis para aquisição, para complementarem os serviços HPE Pointnext Complete Care Starter Pack, são oferecidos blocos de dez(10) crédito. Os clientes podem ampliar o número de créditos necessários para atender às necessidades de sua organização adquirindo vários blocos de dez (10) créditos e os períodos de duração associados a eles."
    table7.cell(2, 0).text = "Dia Técnico HPE"
    table7.cell(2, 1).text = "Esse recurso de serviço opcional é encomendado por quantidade de dias técnicos necessários para o cliente. O serviço Dia Técnico HPE oferece ao cliente a flexibilidade de personalizar as tarefas além do escopo dos serviços técnicos padrão. \nEspecialistas em serviços técnicos altamente treinados podem ajudar o cliente com uma variedade de atividades de operações, otimizações e avaliações. Para mais informações, consulte a Ficha técnica do Dia Técnico HPE. A ficha técnica do Dia Técnico HPE será aplicada a qualquer compra de dias da equipe e é incorporada neste documento como referência."

    p96 = doc.add_paragraph()
    p96.add_run("\n\n3.3.3.	Limitações do Serviço ").bold = True
    p96.add_run("\nAplicavel aos contratos que contenham algum dos Números de produto abaixo:\nH2T12AC; H2T12BC;")
    p96.add_run("\n\nOs serviços HPE Pointnext Complete Care Starter Pack são um conjunto de ofertas de preço fixo que fornecem os recursos de ambiente e gerenciamento de relacionamento do HPE Pointnext Complete Care estabelecidos nesta ficha técnica. \nEsses serviços não incluem nenhum recurso de suporte reativo ou serviços proativos técnicos específicos do produto. O cliente deve adquirir esses serviços separadamente por meio dos serviços do HPE Pointnext Complete Care Add-on e dos Créditos de Serviço HPE, respectivamente.")
    p96.add_run("\n\nOs serviços HPE Pointnext Complete Care foram criados para serem vendidos juntos a ou antes da aquisição do HPE Pointnext Complete Care Add-on. Consulte a seção Limitações do serviço da ficha técnica do HPE Pointnext Complete Care Add-on para ver quais  as ações necessárias quando o período de cobertura da extensão for diferente do período de cobertura do serviço Starter Pack e ver qual o impacto nesse suporte caso essas ações não sejam realizadas.")
    p96.add_run("\n\nO escopo do ambiente dos serviços HPE Pointnext Complete Care Starter Pack é restrito ao ambiente de TI sob o gerenciamento diário direto de uma organização de TI em um país.")
    p96.add_run("\nOs Créditos de Serviço HPE, especificamente as ofertas de 10 créditos, estão disponíveis para aquisição com o serviço HPE Pointnext Complete Care Starter Pack. Se o termo de cobertura dos Créditos de Serviço HPE ultrapassar o prazo de cobertura do serviço Starter Pack, os clientes terão o direito de resgatar seus créditos restantes até o final do termo de duração associado aos créditos de serviço. Entretanto, o recurso da conta designado fornecido sob o serviço HPE Pointnext Complete Care Starter Pack adquirido (seja ASM ou Consultor de Serviço, conforme aplicável) não estará mais disponível para ajudar os clientes a planejarem o uso dos créditos deles. Os clientes podem gerenciar por conta própria os saldos de crédito e selecionar itens no Centro de Suporte HPE. \nOs clientes que precisarem de assistência no planejamento de crédito de um ASM devem resgatar o Serviço de planejamento de crédito no local no menu de Créditos de Serviço.")
    p96.add_run("\n\nOs Créditos de Serviço HPE estão disponíveis apenas para determinados servidores, software, dispositivos de armazenamento, arrays de armazenamento, redes e SANs HPE. Os recursos desse serviço podem variar ou serem limitados com base em dispositivos ou software específicos. Consulte um representante de vendas da HPE para solicitar informações sobre limitações específicas e disponibilidade local.")
    p96.add_run("\n\nOs recursos da conta designados pela HPE fornecem os produtos proativos necessários durante os dias úteis e horários comerciais padrão da HPE, excluindo os feriados da HPE, remotamente ou no local, a critério da HPE. O fornecimento de suporte proativo fora do horário comercial padrão da HPE em dias úteis padrão pode ser adquirido separadamente e está sujeito à disponibilidade local.")

elif ui.servico_final.get() == 2: 

    title94 = doc.add_heading(level=0)
    title94.add_run("3.	Especificações dos Níveis de Serviços")
    title94.add_run("\n3.4.	HPE Pointnext Complete Care Add-on ")
    p97 = doc.add_paragraph()
    p97.add_run("\n\nAplicavel aos contratos que contenham algum dos Números de produto abaixo:\nH2T05AC; H0GT7AC; HA158AC; HA156AC; HL935AC; HA162AC; HA167AC; H7G18AC; H0JC1AC; HA360AC; H2T05AC; HU4A2AC")
    p97.add_run("\n\n3.4.1.	Destaques dos Recursos do Serviço ").bold = True
    p97.add_run("\nTABELA 1. ")
    p97.add_run("Recursos Proativos\n")
    table8 = doc.add_table(rows=1, cols=2)
    table8.style = 'Table Grid'
    table8.autofit = True
    table8.allow_autofit = True
    for row8 in table8.rows:
        for cell8 in row8.cells:
            cell8.width = Pt(180)
    table8.cell(0, 0).text = "\nSujeito a quaisquer limitações, conforme definidas neste documento, o serviço HPE Pointnext Complete Care Add-on estende os seguintes recursos proativos para o novo produto conforme descrito mais especificamente no contrato(SOW) do cliente ou na ficha técnica do serviço HPE Pointnext Complete Care Starter Pack.\nEsses recursos são descritos em um nível alto. Consulte seu contrato ou ficha técnica do HPE Pointnext Complete Care Starter Pack para obter mais detalhes sobre esses recursos de serviço.\n\nO SRM inclui os seguintes itens:\n•	Equipe de conta designada\n•	Gerenciamento da implementação do serviço\n•	Planejamento de suporte à conta\n•	Gerenciamento de inventário\n•	Planejamento e revisão do serviço\n•	Análise de atividade de suporte ou relatório de incidente\n•	Consultoria operacional e técnica ou assistência operacional\n•	Acesso ao Centro de Suporte HPE\n•	Assistência e planejamento para contratação de Serviços Educacionais HPE\n"
    table8.cell(0, 1).text = "\nRecursos proativos para todo ambiente de IT, que podem incluir:\n•	Acesso multicanal a especialistas que conhecem o cliente\n•	Resposta remota 24x7\n•	Acesso e roteamento de incidents prioritário\n•	Ligação direta com a equipe de conta HPE designada\n•	Gerenciamento de escalação acelerada\n•	Experiência digital personalizada\n•	Análise e relatório de incidente\n•	Opções de cobertura expandida\n•	Gerenciamento de chamada global (opcional)\n"
    p98 = doc.add_paragraph()
    p98.add_run("\n\nTABELA 2. ").bold = True
    p98.add_run("Recursos proativos específicos de produto\n")
    table9 = doc.add_table(rows=1, cols=2)
    table9.style = 'Table Grid'
    table9.autofit = True
    table9.allow_autofit = True
    for row9 in table9.rows:
        for cell9 in row9.cells:
            cell9.width = Pt(180)
    table9.cell(0, 0).text = "\nSujeito a quaisquer limitações, conforme definidas neste documento, o serviço de HPE Pointnext Complete Care Add-on pode estender os seguintes recursos proativos específicos do produto para o novo produto se esses recursos específicos do produto estiverem incluídos no seu contrato personalizado (SOW) ou se forem adquiridos por meio do uso dos Créditos de Serviço HPE.\nNota: Os serviços proativos não estão disponíveis em todos os produtos. A equipe da conta designada do HPE Pointnext Complete Care trabalhará com o cliente para determinar quais atividades proativas são as melhores para os produtos no ambiente do cliente.\nServiços de computação, que podem incluir o seguinte:\n•	Análise e gerenciamento de firmware e software de servidor\n•	Verificação de integridade do sistema\n•	Recomendações de implementação – computação\nServiços de armazenamento, que podem incluir o seguinte:\n•	Avaliação da capacidade de suporte da SAN\n•	Análise e gerenciamento de firmware e software de armazenamento/SAN\n•	Avaliação técnica da alta disponibilidade do armazenamento\n•	Recomendações de implementação – armazenamento\n"
    table9.cell(0, 1).text = "\nServiços de rede, que podem incluir o seguinte:\n•	Análise e gerenciamento de firmware e software da rede\n•	Recomendações de implementação – rede\nServiços de ambiente\n•	Créditos de Serviço HPE\n•	Créditos de Treinamento HPE\n•	Dia da Equipe HPE\n•	Dia do engenheiro designado ao cliente (ACE, assigned customer engineer) (disponível somente com o nível de experiência padrão)\n•	Gerenciamento de projetos (disponível somente com o nível de experiência padrão)\nConsulte o seus contrato personalizado (SOW) do HPE Pointnext Complete Care para informações sobre os recursos proativos específicos do produto ou fale com a equipe de conta designada sobre os recursos proativos específicos do produto disponíveis para aquisição usando-se os Créditos de Serviço HPE ou o Dia da Equipe HPE.\n"
    p99 = doc.add_paragraph()
    p99.add_run("\n\nTABELA 3. ").bold = True
    p99.add_run("Recursos Reativos\n")
    table10 = doc.add_table(rows=2, cols=2)
    table10.style = 'Table Grid'
    table10.autofit = True
    table10.allow_autofit = True
    for row10 in table10.rows:
        for cell10 in row10.cells:
            cell10.width = Pt(180)
    table10.cell(0, 0).text = "\nO serviço de HPE Pointnext Complete Care Add-on inclui o suporte reativo HPE Pointnext Tech Care. Alguns recursos de destaque desse suporte são:\nOpções de suporte reativo de hardware HPE Pointnext Tech Care em três níveis\n•	HPE Pointnext Complete Care Add-on Básico\n•	HPE Pointnext Complete Care Add-on Essencial\n•	HPE Pointnext Complete Care Add-on Crítico\nSuporte reativo de software:\n•	Suporte de software 24x7\n"
    table10.cell(0, 1).text = "\nRecursos do serviço de software\n•	Licença para usar atualizações de software se adquiridas da HPE\n•	Atualizações de produto de software e documentação\n•	Suporte consultivo para instalação\nRecursos opcionais\n•	Retenção de mídia defeituosa\n•	Retenção abrangente de material defeituoso\n"
    table10.cell(1, 0).text = "\nPara mais informações relacionadas ao suporte reativo incluído, com os recursos disponíveis e as opções de nível de serviço, consulte a Ficha técnica do HPE Pointnext Tech Care.\n"
    a = table10.cell(1, 0)
    b = table10.cell(1, 1)
    a.merge(b)
    p90 = doc.add_paragraph()
    p90.add_run("\n\n3.4.2.	Limitações do Serviço ").bold = True
    p90.add_run("\n\nO serviço HPE Pointnext Complete Care Add-on é uma oferta de preço fixo disponível na hora da aquisição do produto pelo cliente. As extensões disponíveis têm preço separado do hardware e devem ser adquiridas no mesmo pedido com o hardware. Quando adquiridos com os produtos de hardware, os produtos receberão a cobertura dos serviços de suporte de hardware (prazo fixo de 3, 4 ou 5 anos), com base na opção adquirida.")
    p90.add_run("\n\nO cliente deve ter um contrato ativo (SOW) ativa dos serviços HPE Pointnext Complete Care ou do serviço HPE Pointnext Complete Care Starter Pack   estabelecida para o ambiente onde os novos produtos vão residir com a HPE ou um revendedor HPE autorizado para ser elegível para      comprar qualquer oferta de HPE Pointnext Complete Care Add-on. Além disso, as restrições e limitações a seguir se aplicam a essas extensões conforme detalhado abaixo. O cliente deve levar em consideração essas condições e limitações ao determinar a duração da cobertura de suporte para suas compras de extensões HPE Pointnext Complete Care.")
    p90.add_run("\n\nSe o período da cobertura do  HPE Pointnext Complete Care Add-on para qualquer produto terminar antes do contrato (SOW) do HPE Pointnext Complete Care ou do serviço HPE Pointnext Complete Care Starter Pack, o cliente precisará renovar os produtos do serviço HPE Pointnext Complete Care Add-on para que a duração de tal contrato termine juntamente à duração do Contrato Principal seja contrato (SOW) do HPE Pointnext Complete Care ou do serviço HPE Pointnext Complete Care Starter Pack, conforme aplicável. Isso é necessário para ajudar a garantir suporte ininterrupto e evitar o retorno às cobranças de suporte. Isso pode ser feito por meio do processo de ajuste ou renovação do HPE Pointnext Complete Care. Para mais informações, entre em contato com um representante HPE ou revendedor HPE autorizado.")
    p90.add_run("\n\nSe quaisquer termos de cobertura do HPE Pointnext Complete Care ultrapassarem o contrato (SOW) do HPE Pointnext Complete Care ou do serviço HPE Pointnext Complete Care Starter Pack e não for renovada sob a cobertura do HPE Pointnext Complete Care antes do vencimento, os recursos proativos do HPE Pointnext Complete Care fornecidos com a aquisição do HPE Pointnext Complete Care Add-on, conforme descrito na Tabela 1, não serão mais estendidos a tais produtos e cessarão após o término do contrato HPE Pointnext Complete Care. Nesse momento, você receberá apenas a cobertura de suporte ao produto reativa adquirida ao longo do termo de cobertura da HPE Pointnext Complete Care Add-on ou HPE Pointnext Complete Care Starter Pack. Nenhum reembolso será fornecido para os recursos de serviço reduzidos como resultado da não renovação do contrato ou do serviço Starter Pack.")
    p90.add_run("\n\n3.4.3.	Pré-Requisitos do Serviço ").bold = True
    p90.add_run("\n\nPara ser elegível ao serviço de HPE Pointnext Complete Care Add-on, o cliente já deve ter um contrato do HPE Pointnext Complete Care com a HPE ou um revendedor autorizado da HPE conforme mencionado em Limitações de serviço.")
    p90.add_run("\n\n3.4.4.	Disposições Gerais ").bold = True
    p90.add_run("\n\nConsulte as disposições gerais listadas no contrato (SOW) ou, se estiver comprando com o HPE Pointnext Complete Care Starter Pack, consulte a seção Provisões gerais da ficha técnica do HPE Pointnext Complete Care, que é incorporada a este documento como referência.")
    p90.add_run("\n\nO cliente reconhece e concorda que a HPE pode usar recursos fora do país da aquisição, para fornecimento desses serviços, a menos que seja especificado de outra forma como parte da descrição do recurso do serviço.")


if ui.mv_final.get() == 0:

    title10 = doc.add_heading(level=0)
    title10.add_run("4.	Suporte a Produtos Multivendor outras marcas")
    title10.add_run("\n\n4.1 Detalhes deste Serviço")
    p10 = doc.add_paragraph()
    p10.add_run("\nAplicavel aos contratos que contenham algum dos Números de produto abaixo:\nH7J32AC, H7J34AC, H7J36AC")
    p10.add_run("\nOs serviços de suporte de vários fornecedores HPE fornecem diagnóstico de produto, suporte de hardware no local e suporte de software para produtos selecionados. Você tem a flexibilidade de escolher tempos de resposta e opções de janela de cobertura para atender às suas necessidades específicas de suporte.\n•	Diagnóstico e suporte remotos de problemas \n•	Suporte de hardware no local \n•	Peças e materiais de reposição \n•	Tempo de resposta no local para suporte de hardware")
    p10.add_run("\n\n4.1.2 Tempo de Resposta").bold = True
    p10.add_run("\n\nÉ o prazo compreendido entre o horário de abertura do chamado na Central de Atendimento da HEWLETT PACKARD ENTERPRISE  Brasil até a chegada do técnico ao local, dentro da janela de cobertura.\n•	9x5 – Next Business Day: Cobertura de peças, mão-de-obra e atendimento on-site, 9 horas por dia (das 8:00 às 17:00 horas), 5 dias úteis por semana, excluindo-se os feriados, com tempo de resposta no local até o dia útil seguinte. H7J32AC\n•	24x7 – 4 horas de Resposta: Cobertura de peças, mão-de-obra e atendimento on-site, 24 horas por dia, 7 dias úteis por semana, incluindo feriados, com tempo de resposta no local em até 4 horas. H7J34AC")
    p10.add_run("\n4.1.3 Tempo de Solução / Call-to-Repair\n").bold = True
    p10.add_run("\n\nO serviço assegura acesso direto a uma equipe de suporte especializada, disponível 24 horas por dia 7 dias por semana, que imediatamente começa a estudar o problema. A HEWLETT PACKARD ENTERPRISE  assume o compromisso de restaurar o hardware em no máximo 6 horas contado do momento em que a chamada é registrada na Central de Atendimento ao Cliente da HEWLETT PACKARD ENTERPRISE  até o momento que o produto seja restaurado o seu nível normal de funcionalidade de hardware. \n•	24x7 – 6 horas Call-to-Repair: Cobertura de peças, mão-de-obra e atendimento on-site, 24 horas por dia, 7 dias por semana, incluindo feriados, com solução on-site em até 6 horas. H7J36AC")
    p10.add_run("\n\n4.1.4 Limitações Gerais do Serviço ").bold = True
    p10.add_run("\n\n•	HPE reserva-se o direito de encerrar este serviço sem aviso prévio se o serviço continuado não for razoavelmente viável.\n•	A menos que especificamente solicitado e contratado, este serviço é limitado apenas ao reparo de hardware e não inclui atualizações. O  de nenhuma forma (incluindo firmware, software incorporado ao hardware, sistema operacional, aplicativos ou qualquer outro tipo de software) ou suporte de engenharia do fabricante. ")
    p10.add_run("\n•	Este serviço não inclui atualizações de código de software, patches, correções de bugs, produtos em camadas ou quaisquer atualizações de software fornecidas pelo fabricante do produto subjacente (OEM), incluindo material protegido por direitos autorais, ferramentas de diagnóstico e documentos. \n•	A critério da HPE, este serviço é fornecido usando uma combinação de diagnóstico e suporte remotos, serviços prestados no local e outros métodos de prestação de serviços. \n•	HPE se reserva o direito de fornecer o suporte por meio de um parceiro terceirizado que usa processos de qualidade como HPE. \n•	Os tempos de resposta no local podem ser atrasados se o diagnóstico remoto identificar uma peça específica necessária para o reparo do produto e essa peça estiver sob encomenda e/ou não estiver disponível. \n•	A menos que especificamente solicitado e contratado, o gerenciamento técnico de elevação do OEM não estará disponível. ")
    p10.add_run("\n\nAs seguintes atividades estão excluídas deste serviço e são de responsabilidade do Cliente: \n•	Backup, recuperação e suporte do sistema operacional, software aplicativo e dados \n•	Teste operacional de aplicativos ou testes adicionais solicitados ou exigidos pelo Cliente \n•	Suporte para problemas relacionados à rede fora do hardware suportado \n•	Serviços necessários devido à falha do Cliente em incorporar qualquer correção, reparo, patch ou modificação do sistema fornecido ao Cliente pela HPE \n•	Serviços necessários devido à falha do Cliente em tomar medidas preventivas previamente aconselhados pela HPE \n•	Serviços necessários devido a tentativas não autorizadas de terceiros para instalar, reparar, manter ou modificar hardware, firmware ou software \n•	Serviços necessários devido ao tratamento ou uso inadequado dos produtos ou equipamentos\n•	Não estão previstas e incluídas trocas gratuitas tanto da peça quanto da mão de obra de itens considerados consumíveis do produto (baterias, filtros, fitas magnéticas, entre outros) durante a vigência do Contrato. Estes itens devem ser cotados e adquiridos a parte, de acordo com a necessidade e solicitação do cliente.")
    p10.add_run("\n\n4.1 Pré-requisitos deste Serviço").bold = True
    p10.add_run("\n\nUma vez aceita esta proposta, a HEWLETT PACKARD ENTERPRISE  executará o levantamento e validação do parque de equipamentos a serem suportados através de um processo de diligência. O objetivo é analisar riscos, melhorias e divergências em relação aos baselines definidos nesta proposta, para posterior negociação e resolução de pontos críticos para o desenvolvimento nas fases de Transição, Transformação e Ongoing, gerando transparência e clara definição do escopo contratado, trazendo benefícios na prestação dos serviços.\nSe durante o processo de diligência for constatado defeitos pré-existentes ao contrato, a HEWLETT PACKARD ENTERPRISE  e o cliente negociarão o reparo do equipamento na modalidade de chamado faturado")
    p10.add_run("\n\nAtividades a serem realizadas\n•	Realizar o levantamento (físico e/ou lógico) detalhado dos equipamentos atuais.\n•	Levantar os planos e as necessidades de crescimento.\n•	Realizar o levantamento dos processos operacionais atuais.\n•	Realizar o levantamento das políticas de segurança do(a) ")
    p10.add_run(ui.cliente_final.get())
    p10.add_run("\nResponsabilidades da Hewlett Packard Enterprise\n•	Realizar o levantamento (físico e/ou lógico) detalhado dos equipamentos atuais.\n•	Provisionar kit de peças baseado no levantamento detalhado dos equipamentos para atender o SLA contratado.\n•	Produzir um documento com os resultados dos levantamentos em classificação de potenciais riscos, lacunas e recomendações de soluções para minimizar tais itens.")
    p10.add_run("\nResponsabilidades do(a) ")
    p10.add_run(ui.cliente_final.get())
    p10.add_run("\n•	Fornecer as informações necessárias do ambiente atual em tempo hábil.\n•	Validar e aprovar a documentação entregue. \n•	Liberação de acesso ao site/Datacenter onde se encontram os equipamentos\n•	Envio dos logs solicitados durante a fase de diligência e ao longo do período de contrato para a análise e prestação do serviço de suporte. O atraso e/ou o não envio dos logs solicitados, pode acarretar no não cumprimento dos níveis de serviço acordado.")

# pag 30

if ui.eosl_final.get() == 0:

    title11 = doc.add_heading(level=0)
    title11.add_run("5.	Suporte para produtos HPE que atingiram a data de fim da Vida Útil – EoSL ")
    title11.add_run("\n\n5.1	Detalhes deste Serviço")
    p11 = doc.add_paragraph()
    p11.add_run("\n\nAplicavel aos contratos que contenham algum dos Números de produto abaixo:\nH7J32AC, H7J34AC, H7J36AC")
    p11.add_run("\nA seguir, descrevem-se o serviço de suporte reativo fornecido pela HPE para produtos de vários fornecedores (dispositivos HPE que atingiram a data de “Fim da Vida Útil – EoSL”), fim do período de suporte da HPE e/ou dispositivos de terceiros) incluídos nesta proposta.\nAs solicitações de suporte para equipamentos cobertos por este serviço serão tratadas na Central de Atendimento ao Cliente da HEWLETT PACKARD ENTERPRISE. \n\nEsses serviços são oferecidos como uma conveniência para o cliente e não são uma continuação da garantia padrão da HPE ou dos serviços de suporte da HPE.")
    p11.add_run("\n\n5.1.4	Tempo de Resposta").bold = True
    p11.add_run("\n\nÉ o prazo compreendido entre o horário de abertura do chamado na Central de Atendimento da HEWLETT PACKARD ENTERPRISE  Brasil até a chegada do técnico ao local, dentro da janela de cobertura.\n•	9x5 – Next Business Day: Cobertura de peças, mão-de-obra e atendimento on-site, 9 horas por dia (das 8:00 às 17:00 horas), 5 dias úteis por semana, excluindo-se os feriados, com tempo de resposta no local até o dia útil seguinte. H7J32AC\n•	24x7 – 4 horas de Resposta: Cobertura de peças, mão-de-obra e atendimento on-site, 24 horas por dia, 7 dias úteis por semana, incluindo feriados, com tempo de resposta no local em até 4 horas. H7J34AC")
    p11.add_run("\n\n5.1.5	Tempo de Solução / Call-to-Repair").bold = True
    p11.add_run("\n\nO serviço assegura acesso direto a uma equipe de suporte especializada, disponível 24 horas por dia 7 dias por semana, que imediatamente começa a estudar o problema. A HEWLETT PACKARD ENTERPRISE  assume o compromisso de restaurar o hardware em no máximo 6 horas contado do momento em que a chamada é registrada na Central de Atendimento ao Cliente da HEWLETT PACKARD ENTERPRISE  até o momento que o produto seja restaurado o seu nível normal de funcionalidade de hardware. \n•	24x7 – 6 horas Call-to-Repair: Cobertura de peças, mão-de-obra e atendimento on-site, 24 horas por dia, 7 dias por semana, incluindo feriados, com solução on-site em até 6 horas. H7J36AC")
    p11.add_run("\n\n5.1.6	Limitações Gerais do Serviço ").bold = True
    p11.add_run("\n\n•	Os serviços são limitados apenas ao reparo de hardware. As atualizações de firmware e software da HPE não estão mais disponíveis para esses produtos. \n•	Devido à idade do equipamento envolvido e seu status HPE EOSL, as peças de reposição podem não ser novas e podem ser provenientes de equipamentos usados. \n•	Os tempos de resposta e solução não são garantidos e podem variar de acordo com a disponibilidade de peças de reposição. \n•	A HPE reserva-se o direito de cancelar este serviço se as peças de reposição não estiverem mais disponíveis.")
    p11.add_run("\n•	Devido ao seu status como EOSL da HPE, suporte de engenharia e elevação não estão disponíveis para produtos HPE EOSL; as peças de reposição são fornecidas no estado em que se encontram; e nenhuma garantia de interoperabilidade é fornecida. \n•	Os recursos de monitoramento remoto e ferramentas podem não estar disponíveis para produtos HPE EOSL.\n•	Para os novos contratos, a HPE propõe um SLA de 30 dias para execução do processo de diligência e preparação dos processo operacionais para o atendimento do contrato de suporte.\n•	Em caso não realização do processo de diligência por pendência do cliente, não serão garantidos os cumprimentos dos níveis de serviços acordados.\n•	A HPE propõe uma reunião de kick-off para alinhamento das expectativas, apresentação dos Gerentes de Serviços.\n•	Não estão previstas e incluídas trocas gratuitas tanto da peça quanto da mão de obra de itens considerados consumíveis do produto (baterias, filtros, fitas magnéticas, entre outros) durante a vigência do Contrato. Estes itens devem ser cotados e adquiridos a parte, de acordo com a necessidade e solicitação do cliente.")
    doc.add_page_break()

# pag 32

title12 = doc.add_heading(level=0)
title12.add_run("Condições Comerciais")
p12 = doc.add_paragraph()
p12.add_run("\nConfiguração e Preços").bold = True
p12.add_run("\nOs produtos (Hardware e Software) e seus respectivos níveis de serviço encontram-se detalhados no Anexo de Configuração e Preços.  (Opcional a critério de Vendas)")

# adicionar tabela

p13 = doc.add_paragraph()
p13.add_run("\n\nO valor total desta proposta é de R$ (informar valor).")
# adicionar valor da proposta? perguntar ao Weid depois
p13.add_run("\n\nAjustes no calendário de faturamento (valor mensal), são aplicáveis desde que não haja alteração no valor total desta proposta e esteja em acordo com a cobertura de suporte e/ou garantia que cada equipamento possuir, além do firmado nesta proposta. De qualquer forma, toda e qualquer mudança no calendário de faturamento dependerá de aprovação prévia da HEWLETT PACKARD ENTERPRISE .")
p13.add_run("\n\nAs condições comerciais aqui estabelecidas foram baseadas nas configurações e distribuição geográfica dos equipamentos, informados pelo cliente.  Quaisquer alterações, as quais deverão ser informadas previamente, com 30 (trinta) dias de antecedência, resultarão na revalidação das condições comerciais, técnicas e tributárias incluindo possível impacto no atendimento ao nível de serviço acordado e aumento dos preços originalmente contratados.")
p13.add_run("\n\nOs preços, condições e valores desta proposta consideram que os equipamentos objeto dos serviços encontram-se em bom estado de conservação e em sua configuração original. A HPE durante o prazo de 30 (trinta) dias, contados do início dos serviços, efetuará a análise técnica, presencial ou remota, do parque de equipamentos objeto do contrato a fim de verificar se os equipamentos se encontram nas condições mínimas de prestação de serviços. Caso a HPE identifique que os equipamentos não são elegíveis aos serviços contratados, enviará ao cliente os custos dos reparos necessários aos equipamentos; e, caso o cliente não efetue o respectivo pagamento ou não autorize os serviços, os equipamentos serão automaticamente excluídos dos serviços. Ainda, durante a vigência dos serviços, caso a HPE identifique que ocorreram alterações das configurações do equipamento, incluindo suas peças, os mesmos serão (a) automaticamente excluídos dos serviços; ou (b) mantidos sob o escopo contratual, desde que o Cliente efetue o pagamento dos valores necessários para o reestabelecimento do equipamento.")
p13.add_run("\n\nNão estão previstas e incluídas trocas gratuitas tanto da peça quanto da mão de obra de itens considerados consumíveis do produto (baterias, filtros, fitas magnéticas, entre outros) durante a vigência do Contrato. Estes itens devem ser cotados e adquiridos a parte, de acordo com a necessidade e solicitação do cliente. A substituição destes itens deve ser feita por produtos originais e compatíveis. A não utilização de produtos originais e compatíveis que venham a causar dano ao produto, exclui o mesmo do contrato até a sua adequação, eliminando-se o equipamento do suporte.")
doc.add_page_break()

# pag 33

title13 = doc.add_heading(level=0)
title13.add_run("Condições de Pagamento\n\n")

if ui.pagamento_final.get() == 0:
    p14 = doc.add_paragraph()
    p14.add_run("A parcela mensal é devida no primeiro dia útil de cada mês de vigência do presente CONTRATO e deverá ser paga à HEWLETT PACKARD ENTERPRISE  em sua sede administrativa (ou a quem e onde a HEWLETT PACKARD ENTERPRISE  indicar), até o último dia útil do mesmo mês ou até a data de vencimento fixada na respectiva fatura. \n")
    doc.add_page_break()
elif ui.pagamento_final.get() == 1:
    p15 = doc.add_paragraph()
    p15.add_run("A parcela à vista é devida no primeiro dia útil do primeiro mês de vigência do presente CONTRATO e deverá ser paga à HEWLETT PACKARD ENTERPRISE  em sua sede administrativa (ou a quem e onde a HEWLETT PACKARD ENTERPRISE  indicar), até o último dia útil do mesmo mês ou até a data de vencimento fixada na respectiva fatura.\n")
    doc.add_page_break()
elif ui.pagamento_final.get() == 2:
    p16 = doc.add_paragraph()
    p16.add_run("A parcela trimestral é devida no primeiro dia útil de cada trimestre de vigência do presente CONTRATO e deverá ser paga à HEWLETT PACKARD ENTERPRISE  em sua sede administrativa (ou a quem e onde a HEWLETT PACKARD ENTERPRISE  indicar), até o último dia útil do mesmo mês ou até a data de vencimento fixada na respectiva fatura.\n")
    doc.add_page_break()

# pag 34

title14 = doc.add_heading(level=0)
title14.add_run("Tributos")
p17 = doc.add_paragraph()
p17.add_run("\n\nOs valores apresentados acima incluem todos os impostos e tributos incidentes sobre os serviços ou produtos objeto desta Proposta de acordo com a legislação em vigor aplicáveis ao local de prestação dos serviços e entrega de produtos, exceto ICMS diferencial de alíquota.")
p17.add_run("\n\nNa eventualidade de ser devido diferencial de aliquota do ICMS ao Estado de destino, o pagamento do valor correspondente será de acordo com a legislação vigente. ")
p17.add_run("\n\nCada espécie ou tipo de serviço ou produto previsto na presente Proposta será faturado de forma segregada a fim de atender a legislação do local da prestação dos serviços ou entrega de produtos, dessa forma, a Hewlett Packard Enterprise emitirá notas-fiscais específicas para cada serviço ou produto fornecido. O estabelecimento que emitir a fatura de serviços será o mesmo do prestador desse serviço, objeto da presente proposta. ")
p17.add_run("\n\nOs serviços a serem prestados não constituem de qualquer forma cessão de mão-de-obra e não contemplam quaisquer retenções de tributos, impostos, contribuições ou taxas pelo Cliente, dessa forma, caso o Cliente efetue qualquer retenção, deverá informar previamente a Hewlett Packard Enterprise, bem como as condições e valores ora ofertados serão revistos diante do impacto ocasionado por eventual retenção. Ainda, o Cliente deverá enviar à Hewlett Packard Enterprise todos os comprovantes e guias de pagamentos da referida retenção.")
p17.add_run("\n\nQuaisquer tributos ou encargos criados, alterados ou extintos, bem como a superveniência de disposições legais quando ocorridas após a data de apresentação da Proposta, e de comprovada repercussão nos preços ofertados, implicarão na revisão destes para mais ou para menos, conforme o caso.")
doc.add_page_break()

# pag 35

title15 = doc.add_heading(level=0)
title15.add_run("Condições para Aceitação do Pedido e Faturamento")
p18 = doc.add_paragraph()
p18.add_run("\n\nSujeitas à aprovação de crédito. \n\nPara clientes novos e/ou situação de renovação,solicitamos os seguintes documentos para análise:\n•	Cópia do Cartão de CNPJ\n•	Ficha cadastral atualizada\n•	Contrato Social e/ou Estatuto Social consolidado e/ou últimas alterações\n•	Relação de Faturamento de, no mínimo, dois últimos Exercícios.\n•	Balanços Patrimoniais e Demonstrações de Resultado dos 03 últimos Exercicios, preferencialmente auditados\n•	Notas explicativas de Balanço\n•	Balancete dos últimos 06 (seis) meses")
p18.add_run("\n\nEstado de Faturamento: ")
p18.add_run(ui.estado_final.get())
p18.add_run("\nO faturamento poderá ser efetuado por uma das filiais da HEWLETT PACKARD ENTERPRISE  abaixo relacionadas.")
p18.add_run("\n\nInício dos Serviços").bold = True
p18.add_run("\n\nA vigência dos serviços dar-se-á de acordo com a data estabelecida no item Configuração e Preços desta proposta. \n\nConsiderando os trâmites administrativos para ativação da nova vigência, solicitamos que a presente proposta seja retornada para HEWLETT PACKARD ENTERPRISE  devidamente assinada em até 05 (cinco) dias úteis anteriores ao término da vigência do contrato ora renovado. O atraso deste retorno, que não poderá ser superior a 05 (cinco) dias úteis, poderá gerar o atraso proporcional no cadastro do novo período contratual, com eventuais reflexos na prestação de serviços de suporte contratados.")
if ui.renovacao_final.get() == 1:
    p19 = doc.add_paragraph()
    p19.add_run("\n\nFindo o prazo de vigência ali indicado, o contrato será automaticamente renovado por iguais e sucessivos períodos, salvo expressa manifestação em contrário da parte interessada, com ao menos 30 (trinta) dias de antecedência ao término do período contratual então em vigor.")
p20 = doc.add_paragraph()
p20.add_run("\n\nConsiderando a necessidade de disponibilização logística dos kits para atendimento e contrato de parcerias, o compromisso de Tempo de Solução / Call-to-Repair, bem como suporte a equipamentos de outros fornecedores (Multivendor), iniciará após 30 dias a partir da assinatura desta proposta, todavia, durante este período, a HEWLETT PACKARD ENTERPRISE envidará todos os esforços para cumprimento dos níveis de serviço detalhados no item configuração e preços da proposta.")
p20.add_run("\n\nReajuste de Preços").bold = True
p20.add_run("\nEm caso de prorrogação contratual, a parcela mensal será reajustada a cada 12 (doze) meses a contar da data da apresentação da proposta, respeitando o interregno de 12 meses, pela variação do IGP-M (Índice Geral de Preços de Mercado, medido pela Fundação Getúlio Vargas), tendo como base o mês anterior ao da emissão do Anexo de Configuração e Preços, do Pedido de Compra ou da aceitação eletrônica, o que tiver ocorrido primeiro.\nCaso o índice no período esteja negativo, não haverá redução nos valores contratuais.")
p20.add_run("\n\nPenalidades por não Cumprimento do Prazo de Pagamento").bold = True
p20.add_run("\n\nEm caso de atraso no pagamento de qualquer quantia devida em decorrência dos Produtos e Serviços prestados, incluindo reembolso de despesas, referidos valores serão acrescidos de multa moratória no importe de 2% (dois por cento), juros de mora de 1% (um por cento) ao mês e correção monetária a ser calculada pela variação do IGP-M (FGV), estes últimos calculados pro rata temporis.")
p20.add_run("\n\nNa hipótese de ser aceito pela HEWLETT PACKARD ENTERPRISE  o envio mensal, pelo cliente, de Pedidos de Compra relacionados a este contrato, o primeiro Pedido de Compra deverá ser encaminhado em até 30 dias da data de aceite desta proposta e os demais deverão ser encaminhados, sucessivamente, no prazo de 30 dias. Eventual atraso no recebimento de um Pedido de Compra, impactando o prazo de pagamento, ensejará a imediata aplicação da referida penalidade. ")
doc.add_page_break()

# Pag 37

title16 = doc.add_heading(level=0)
title16.add_run("Validade da Proposta")
p23 = doc.add_paragraph()
p23.add_run("\nVálida até ")
p23.add_run(ui.str_vl)
p23.add_run("\n\nDados Cadastrais\n").bold = True
table11 = doc.add_table(rows=4, cols=1)
table11.style = 'Table Grid'
table11.autofit = False
table11.allow_autofit = False
for row11 in table11.rows:
    for cell11 in row11.cells:
        cell11.width = Pt(180)
table11.cell(0, 0).text = "Para Efeito de Faturamento"
table11.cell(1, 0).text = "\nFilial Barueri\nHewlett-Packard Brasil Ltda. \nEndereço: Alameda Rio Negro, nº. 750, Térreo, Sala Rio de Janeiro\nCEP: 06454-000 – Alphaville  – Barueri – SP\nCNPJ: 61.797.924/0002-36\nInscrição Estadual: 206.203.581.118\nInscrição Municipal: 5.39.202-6\n"
table11.cell(2, 0).text = "\nFilial Porto Alegre\nHewlett-Packard Brasil Ltda \nEndereço: Av. Ipiranga, 6681 – Prédio 91B – Azenha \nCEP: 90610-001 – Porto Alegre/ RS \nCNPJ: 61.797.924/0013-99 \nInscrição Estadual: 096/2962120\nInscr. Municipal: 20409524\n"
table11.cell(3, 0).text = "\nFilial Rio de Janeiro\nHewlett-Packard Brasil Ltda\nEndereço: Avenida Almirante Barroso, 00081, SAL 33B113 PARTE 2, Centro\nCEP: 20031-004 - Rio de Janeiro – RJ \nCNPJ: 61.797.924/0016-31 \nInscrição Estadual: 77.496.980\nInscr. Municipal: 0.331.073-6\n"
p24 = doc.add_paragraph()
p24.add_run("\n\n\n")
table12 = doc.add_table(rows=4, cols=2)
table12.style = 'Table Grid'
table12.autofit = True
table12.allow_autofit = True
for row12 in table12.rows:
    for cell12 in row12.cells:
        cell12.width = Pt(180)
table12.cell(0, 0).text = "Descrição do serviço"
table12.cell(0, 1).text = "Código de serviço da NF"
table12.cell(1, 0).text = "Manutenção de equipamento para informática e automação"
table12.cell(1, 1).text = "14.01.13214"
table12.cell(2, 0).text = "Treinamento"
table12.cell(2, 1).text = "08.02.14211"
table12.cell(3, 0).text = "Software"
table12.cell(3, 1).text = "01.07.01217"
p25 = doc.add_paragraph()
p25.add_run("\n\nOs produtos e serviços objeto da presente proposta poderão ser prestados dos estabelecimentos da Hewlett Packard Enterprise, e a Hewlett Packard Enterprise poderá, a seu exclusivo critério, executar os Serviços ou qualquer porção dos Serviços por meio de uma Afiliada da Hewlett Packard Enterprise ou empresa pertencente ao seu Grupo Econômico. As obrigações da Hewlett Packard Enterprise sob esta Proposta e futuro Contrato serão executadas pela própria Hewlett Packard Enterprise ou por suas subsidiárias ou afiliadas, sem qualquer limitação.")
doc.add_page_break()

# pag 38

if ui.proposta_final.get() == 0:

    title17 = doc.add_heading(level=0)
    title17.add_run("5.	Condições Gerais")
    title17.add_run("\n\nObrigações da Hewlett Packard Enterprise")
    p26 = doc.add_paragraph()
    p26.add_run("\n\nA HEWLETT PACKARD ENTERPRISE  responsabiliza-se por todos os ônus e encargos trabalhistas e previdenciários resultantes da contratação e emprego de pessoas para a realização dos serviços objeto desta proposta, bem como pelos excessos e omissões praticados pelos mesmos, razão pela qual é a HEWLETT PACKARD ENTERPRISE  considerada empregadora autônoma, não existindo entre ela e seus empregados, vínculo de qualquer natureza com o Cliente.")
    p26.add_run("\n\nA HEWLETT PACKARD ENTERPRISE  obriga-se a não divulgar, reproduzir, vender ou utilizar a favor de terceiros, os trabalhos e demais dados obtidos no Cliente, através da prestação dos serviços descriminados nesta proposta, sob a pena de responsabilidade civil e criminal.")
    p26.add_run("\n\nA HEWLETT PACKARD ENTERPRISE  obriga-se, ainda, a executar os serviços, dentro das boas técnicas e dos costumes usuais em trabalho deste gênero, bem como a utilização de mão-de-obra qualificada.")
    p26.add_run("\n\nTodos os tributos (impostos, taxas e contribuições) de natureza federal, estadual e municipal, incidentes ou que venham a incidir sobre o objeto deste contrato serão de responsabilidade única da HEWLETT PACKARD ENTERPRISE .")


    p27 = doc.add_paragraph()
    p27.add_run("\n\nRecisão").bold = True
    p27.add_run("\n\nQualquer das partes poderá rescindir este Contrato total ou parcialmente, mediante notificação por escrito à outra parte com 90 (noventa) dias de antecedência, exceto nos casos em que os equipamentos se tornarem obsoletos, hipótese em que este Contrato pode ser rescindido total ou parcialmente mediante notificação com 30 (trinta) dias de antecedência.  ")

    if ui.rescisao_final.get() == 0:
        p27.add_run("\nRescisão atual: 90 dias")
    elif ui.rescisao_final.get() == 1:
        p27.add_run("\nRescisão atual: 30 dias")

    p28 = doc.add_paragraph()
    p28.add_run("\n\nTransferência").bold = True
    p28.add_run("\n\nEste acordo, bem como qualquer direito ou obrigação aqui estabelecidos, não poderão ser cedidos ou transferidos, no todo ou em parte, por qualquer uma das partes sem o prévio consentimento por escrito da outra. No entanto, a HEWLETT PACKARD ENTERPRISE  terá o direito de ceder ou transferir quaisquer dos seus direitos e obrigações decorrentes do presente contrato, caso ocorra qualquer alteração na estrutura societária da HEWLETT PACKARD ENTERPRISE , parcial ou total (por “spin-off”, cisão, alienação de bens ou mediante qualquer outra operação societária semelhante). A HEWLETT PACKARD ENTERPRISE  deverá notificar o Cliente, por escrito e em tempo razoável, após qualquer cessão ou transferência de quaisquer de suas obrigações ou direitos, conforme disposto acima. ")
    p28.add_run("\n\nTermos e Condições Contratuais").bold = True
    p28.add_run("\n\nAos Produtos e Serviços aqui ofertados aplicam-se o “Portfolio Terms with Supplemental Data Sheet Terms for Support (CTPF01 and CTDS01)”, registrado no Cartório Oficial de Registro de Títulos e Documentos da Comarca de Barueri – SP, protocolizado em títulos e documentos sob o nº 1774062 em 18 de outubro de 2021.\nCaso haja condições comerciais descritas nessa proposta divergentes das descritas no “Portfolio Terms with Supplemental Data Sheet Terms for Support (CTPF01 and CTDS01)”, prevalecem as condições da proposta.\nO Cliente, ao emitir um Pedido de Compra, formalizando a contratação com base nesta Proposta, aceita que:")
    p28.add_run("\n\n(i) os termos e condições constantes do “Portfolio Terms with Supplemental Data Sheet Terms for Support (CTPF01 and CTDS01)” são os únicos aplicáveis à contratação, excluindo-se expressamente todos os termos e condições padrão do Cliente que constem do referido Pedido, termos e condições estes que não produzirão quaisquer efeitos entre as partes, a qualquer título, mesmo que o Pedido do Cliente seja aceito pela HEWLETT PACKARD ENTERPRISE  e mesmo que haja disposição em contrário no referido Pedido de Compra. A HEWLETT PACKARD ENTERPRISE , ao apresentar esta Proposta e ao aceitar eventual Pedido de Compra, rejeita expressamente quaisquer termos e condições padrão do Cliente. \nEstes Termos constituem o acordo integral entre a HEWLETT PACKARD ENTERPRISE  e o Cliente a respeito da compra pelo Cliente de produtos e serviços da HEWLETT PACKARD ENTERPRISE , e substituem e cancelam quaisquer comunicações, declarações ou acordos anteriores, assim como quaisquer condições adicionais ou inconsistentes do Cliente, sejam verbais ou por escrito.")
    p28.add_run("\n\nAcordo Integral").bold = True
    p28.add_run("\n\nAo aceitar a presente Proposta, mediante assinatura deste instrumento e/ou emissão do Pedido de Compra, a mesma passa a ter força de contrato entre as partes.  Os Termos, referidos acima, constituem o acordo integral entre a HEWLETT PACKARD ENTERPRISE  e o Cliente a respeito da compra pelo Cliente de produtos e serviços da HEWLETT PACKARD ENTERPRISE , e substituem e cancelam quaisquer comunicações, declarações ou acordos anteriores, assim como quaisquer condições adicionais ou inconsistentes do Cliente, sejam verbais ou por escrito.")
    p28.add_run("\n\nRestrições de Uso e Divulgação da Proposta").bold = True
    p28.add_run("\n\nAs informações (dados) que constam de todas as folhas deste documento/cotação constituem informações confidenciais da Hewlett Packard Enterprise Ltda. (doravante referenciada “HEWLETT PACKARD ENTERPRISE ”). As informações fornecidas ao cliente não podem ser usadas ou divulgadas, sem a prévia autorização da HEWLETT PACKARD ENTERPRISE , para propósitos que não sejam os de avaliação da proposta.\nAs propostas da HEWLETT PACKARD ENTERPRISE  poderão ser submetidas via email e mídia eletrônica para sua conveniência. Se o conteúdo diferenciar entre as cópias impressas e mídia eletrônica, somente o conteúdo da impressa será garantido pela HEWLETT PACKARD ENTERPRISE.")
    doc.add_page_break()

elif ui.proposta_final.get() == 1:

    title17 = doc.add_heading(level=0)
    title17.add_run("5.	Condições Gerais")
    title17.add_run("Termos e Condições Contratuais")
    p26 = doc.add_paragraph()
    p26.add_run("\n\nAos Produtos e Serviços aqui ofertados aplicam-se o disposto na Lei nº 14.133/2021e demais legislações aplicáveis ao assunto. \n\nCaso haja condições comerciais descritas nessa proposta divergentes das descritas no contrato firmado entre as partes, prevalecem as condições da proposta.")
    p26.add_run("\n\nAcordo Integral").bold = True
    p26.add_run("\n\nEstes Termos constituem o acordo integral entre a HEWLETT PACKARD ENTERPRISE   e o Cliente a respeito da compra pelo Cliente de produtos e serviços da HEWLETT PACKARD ENTERPRISE, e substituem e cancelam quaisquer comunicações, declarações ou acordos anteriores, assim como quaisquer condições adicionais ou inconsistentes do Cliente, sejam verbais ou por escrito.")
    p26.add_run("\n\nRestrições de Uso e Divulgação da Proposta").bold = True
    p26.add_run("\n\nAs informações (dados) que constam de todas as folhas deste documento/cotação constituem informações confidenciais da Hewlett Packard Enterprise Ltda. (doravante referenciada “HEWLETT PACKARD ENTERPRISE  ”). As informações fornecidas ao cliente não podem ser usadas ou divulgadas, sem a prévia autorização da HEWLETT PACKARD ENTERPRISE, para propósitos que não sejam os de avaliação da proposta. \nAs propostas da HEWLETT PACKARD ENTERPRISE   poderão ser submetidas via email e mídia eletrônica para sua conveniência. Se o conteúdo diferenciar entre as cópias impressa e mídia eletrônica, somente o conteúdo da impressa será garantido pela HEWLETT PACKARD ENTERPRISE.")
    p26.add_run("\n\nDeclaração de Anticorrupção").bold = True 
    p26.add_run("\n\nAs partes declaram que seus empregados, representantes, dirigentes ou administradores, direta ou indiretamente, não efetuaram qualquer promessa, oferta, solicitação ou aceite de qualquer vantagem indevida, de qualquer natureza, para a execução desta Proposta ou que esteja sob qualquer forma conexa com a presente Proposta, bem como obrigam-se a cumprir, o disposto na Lei nº. 12.846/2013.")
    p26.add_run("\n\nAs partes declaram, pelo presente, que não praticam ou praticarão qualquer atividade que viole qualquer legislação anticorrupção, incluindo, o US Foreign Corrupt Practices Act – FCPA, o UK Bribery Act, a Lei Federal nº. 12.846/13 e quaisquer Decretos, Leis Estaduais ou Municipais.")
    p26.add_run("\n\nAs partes declaram que têm conhecimento da Lei nº. 12.846/2013, bem com comprometem-se a não praticar qualquer dos atos lesivos à Administração Pública elencados no Art. 5º, incisos e alíneas, seja durante o certame licitatório, seja no decorrer da execução do contrato, sob pena de responsabilização, independente da aferição de culpa. Declara, ainda, para os devidos fins, estar ciente das sanções previstas na referida legislação, além daquelas cominadas na Lei nº. 14.133/2021, demais normas de licitações e contratos da Administração Pública e normas correlatas.")
    p26.add_run("\n\nObrigações da Hewlett Packard Enterprise").bold = True
    p26.add_run("\n\nA HPE responsabiliza-se por todos os ônus e encargos trabalhistas e previdenciários resultantes da contratação e emprego de pessoas para a realização dos serviços objeto desta proposta, bem como pelos excessos e omissões praticados pelos mesmos, razão pela qual é a HPE considerada empregadora autônoma, não existindo entre ela e seus empregados, vínculo de qualquer natureza com o Cliente.")
    p26.add_run("\n\nA HPE obriga-se a não divulgar, reproduzir, vender ou utilizar a favor de terceiros, os trabalhos e demais dados obtidos no Cliente, através da prestação dos serviços descriminados nesta proposta, sob a pena de responsabilidade civil e criminal.\n\nA HPE obriga-se, ainda, a executar os serviços, dentro das boas técnicas e dos costumes usuais em trabalho deste gênero, bem como a utilização de mão-de-obra qualificada.\n\nTodos os tributos (impostos, taxas e contribuições) de natureza federal, estadual e municipal, incidentes ou que venham a incidir sobre o objeto deste contrato serão de responsabilidade única da HPE.")

    p27 = doc.add_paragraph()
    p27.add_run("\n\nRecisão").bold = True
    p27.add_run("\n\nQualquer das partes poderá rescindir este Contrato total ou parcialmente, mediante notificação por escrito à outra parte com 90 (noventa) dias de antecedência, exceto nos casos em que os equipamentos se tornarem obsoletos, hipótese em que este Contrato pode ser rescindido total ou parcialmente mediante notificação com 30 (trinta) dias de antecedência.  ")

    if ui.rescisao_final.get() == 0:
        p27.add_run("\nRescisão atual: 90 dias")
    elif ui.rescisao_final.get() == 1:
        p27.add_run("\nRescisão atual: 30 dias")

    doc.add_page_break()

# pag 40

title18 = doc.add_heading(level=0)
title18.add_run("6.	Termo de Aceite da Proposta / Pedido de Compra")
p29 = doc.add_paragraph()
p29.add_run("A Hewlett Packard Enterprise.")
p29.add_run("\nA/C:  Sr. (a) ")
p29.add_run(ui.vendedor.get())
p29.add_run("/ Tel: ")
p29.add_run(ui.var_telefone.get())
p29.add_run("\nEste documento tem como objetivo formalizar o aceite da proposta/cotação da Hewlett-Packard número OPE:")
p29.add_run(ui.ope_final.get())
p29.add_run(", referente ao projeto")
p29.add_run(ui.num_contrato_final.get())
p29.add_run(", datado de")
p29.add_run(ui.str_dt)
p29.add_run("\n\n")
table13 = doc.add_table(rows=36, cols=3)
table13.style = 'Table Grid'
table13.autofit = True
table13.allow_autofit = True
for row13 in table13.rows:
    for cell13 in row13.cells:
        cell13.width = Pt(250)
table13.cell(0, 0).text = "DECLARAÇÃO DE ACEITE DESTA PROPOSTA E ANEXOS"
a = table13.cell(0, 0)
b = table13.cell(0, 2)
a.merge(b)
table13.cell(1, 0).text = "Ao aceitar esta Proposta declaro estar ciente e de acordo com a proposta e todas as cláusulas e condições do “Portfolio Terms with Supplemental Data Sheet Terms for Support (CTPF01 and CTDS01)”, item Termos e Condições Contratuais desta proposta. "
a = table13.cell(1, 0)
b = table13.cell(1, 2)
a.merge(b)
table13.cell(2, 0).text = "RESPONSÁVEL PELO ACEITE DESTA PROPOSTA"
a = table13.cell(2, 0)
b = table13.cell(2, 2)
a.merge(b)
table13.cell(3, 0).text = "Nome:"
a = table13.cell(3, 0)
b = table13.cell(3, 1)
a.merge(b)
table13.cell(3, 2).text = "Cargo:"
table13.cell(4, 0).text = "Departamento:"
a = table13.cell(4, 0)
b = table13.cell(4, 2)
a.merge(b)
table13.cell(5, 0).text = "Assinatura:"
a = table13.cell(5, 0)
b = table13.cell(5, 2)
a.merge(b)
table13.cell(6, 0).text = "Local e Data:"
a = table13.cell(6, 0)
b = table13.cell(6, 2)
a.merge(b)
table13.cell(7, 0).text = "Razão Social:"
a = table13.cell(7, 0)
b = table13.cell(7, 2)
a.merge(b)
table13.cell(8, 0).text = "INFORMAÇÃO DE FATURAMENTO"
a = table13.cell(8, 0)
b = table13.cell(8, 2)
a.merge(b)
table13.cell(9, 0).text = "Razão Social:"
a = table13.cell(9, 0)
b = table13.cell(9, 2)
a.merge(b)
table13.cell(10, 0).text = "Endereço de Faturamento:"
a = table13.cell(10, 0)
b = table13.cell(10, 2)
a.merge(b)
table13.cell(11, 0).text = "CNPJ:"
a = table13.cell(11, 0)
b = table13.cell(11, 1)
a.merge(b)
table13.cell(12, 0).text = "Inscr. Estadual:"
a = table13.cell(12, 0)
b = table13.cell(12, 2)
a.merge(b)
table13.cell(13, 0).text = "O faturamento da presente contratação não dependerá da emissão de Pedido de Compra e/ou Contrato específico, assim, a presente Proposta será o único instrumento contratual válido entre a HPE e o Cliente. Em caso de emissão de Pedido de Compra e/ou Contrato, solicitamos assinalarem o campo ao lado: (   )\n\n** a emissão de Pedido de Compra será entendida como o aceite integral dessa Proposta, sendo que os termos e condições deste documento prevalecem sobre quaisquer outras disposições anteriormente ou posteriormente emitidas e regem a presente relação comercial entre a HPE e o Cliente.   "
a = table13.cell(13, 0)
b = table13.cell(13, 2)
a.merge(b)
table13.cell(14, 0).text = "Aceita faturamento parcial:     (   ) sim          (    ) não"
a = table13.cell(14, 0)
b = table13.cell(14, 2)
a.merge(b)
table13.cell(15, 0).text = "Aceita diferença de centavos:     (   ) sim          (    ) não"
a = table13.cell(15, 0)
b = table13.cell(15, 2)
a.merge(b)
table13.cell(16, 0).text = "Possui janela de faturamento:     (   ) sim          (    ) não \n\nCaso positivo informar data limite mensal para recebimento das notas fiscais: _____________"
a = table13.cell(16, 0)
b = table13.cell(16, 2)
a.merge(b)
table13.cell(17, 0).text = "INFORMAÇÃO E LOCAL DE ENTREGA"
a = table13.cell(17, 0)
b = table13.cell(17, 2)
a.merge(b)
table13.cell(18, 0).text = "Razão Social:"
a = table13.cell(18, 0)
b = table13.cell(18, 2)
a.merge(b)
table13.cell(19, 0).text = "Endereço de entrega:"
a = table13.cell(19, 0)
b = table13.cell(19, 2)
a.merge(b)
table13.cell(20, 0).text = "CNPJ:"
a = table13.cell(20, 0)
b = table13.cell(20, 1)
a.merge(b)
table13.cell(21, 0).text = "Inscr. Estadual:"
table13.cell(22, 0).text = "Contato:"
a = table13.cell(22, 0)
b = table13.cell(22, 1)
a.merge(b)
table13.cell(23, 0).text = "Telefone:"
table13.cell(24, 0).text = "E-mail para envio da mídia eletrônica de software::"
a = table13.cell(24, 0)
b = table13.cell(24, 2)
a.merge(b)
table13.cell(25, 0).text = "*** (Caso o CNPJ do local de entrega seja diferente do CNPJ da empresa contratante, favor informar se é um operador logístico/filial ou um datacenter e o contato: \n\n"
a = table13.cell(25, 0)
b = table13.cell(25, 2)
a.merge(b)
table13.cell(26, 0).text = "Necessário pré-agendamento? (   ) sim          (    ) não"
a = table13.cell(26, 0)
b = table13.cell(26, 2)
a.merge(b)
table13.cell(27, 0).text = "Possuim restrição de horário para recebimento do produto? (   ) sim          (    ) não\nCaso positivo, informar os horários.:\n\n"
a = table13.cell(27, 0)
b = table13.cell(27, 2)
a.merge(b)
table13.cell(28, 0).text = "INFORMAÇÃO DE COBRANÇA:"
a = table13.cell(28, 0)
b = table13.cell(28, 2)
a.merge(b)
table13.cell(29, 0).text = "Endereço de Cobrança:"
a = table13.cell(29, 0)
b = table13.cell(29, 2)
a.merge(b)
table13.cell(30, 0).text = "CNPJ:"
a = table13.cell(30, 0)
b = table13.cell(30, 1)
a.merge(b)
table13.cell(31, 0).text = "Inscr. Estadual:"
table13.cell(32, 0).text = "Contato:"
a = table13.cell(32, 0)
b = table13.cell(32, 1)
a.merge(b)
table13.cell(33, 0).text = "Telefone:"
table13.cell(34, 0).text = "E-mail para envio da Nota Fiscal Eletrônica:"
a = table13.cell(34, 0)
b = table13.cell(34, 2)
a.merge(b)
table13.cell(35, 0).text = "Comentários Adicionais:"
a = table13.cell(35, 0)
b = table13.cell(35, 2)
a.merge(b)

doc.add_page_break()

title19 = doc.add_heading(level=0)
title19.add_run("\n7.	Anexos ")

# Salvar o documento

doc.save('example_document.docx')