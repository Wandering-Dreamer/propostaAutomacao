import pandas as pd
from openpyxl import load_workbook
from openpyxl.styles import Alignment
from openpyxl.styles import PatternFill
from openpyxl.styles import Border, Side
from openpyxl.utils.dataframe import dataframe_to_rows
from datetime import datetime
from openpyxl.utils import get_column_letter
import csv
from openpyxl import Workbook
import docx
from docx.shared import Pt, Mm, Cm, Inches
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def change_table_cell(cell, background_color=None, font_color=None, font_size=None, bold=None, italic=None):

    if background_color:
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), background_color))
        cell._tc.get_or_add_tcPr().append(shading_elm)

    if font_color:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = docx.shared.RGBColor.from_string(font_color)

    if font_size:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = docx.shared.Pt(font_size)

    if bold is not None:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = bold

    if italic is not None:
        for p in cell.paragraphs:
            for r in p.runs:
                r.italic = italic

def change_table_row(table_row, background_color=None, font_color=None, font_size=None, bold=None, italic=None):
    for cell in table_row.cells:
        change_table_cell(cell, background_color=background_color, font_color=font_color, font_size=font_size,
                          bold=bold,
                          italic=italic)

df = pd.read_excel("SA38 Test.xlsx")
doc = docx.Document()
unique_values = df.iloc[:, 16].drop_duplicates().index
print(unique_values)
merge_df = [2, 3]
text = pd.Series(["Vigência - De:"])
text2 = pd.Series(["Vigência - Até:"])
thin = Side(border_style="thin", color="000000")

section = doc.sections[0]

section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

writer = pd.ExcelWriter("pandas_to_excel.xlsx") 
for i in unique_values:
    if i == 1:
        x = 2

        fl = df.iloc[i:(unique_values[x]), 16]
        fl = fl.drop_duplicates()
        fl.to_excel(writer, sheet_name='sheetName', index=False, startrow=1, startcol=0, header=False)
        sla = df.iloc[i-1:(unique_values[x]), 9]
        sla = sla.drop_duplicates()
        sla.to_excel(writer, sheet_name='sheetName', index=False, startrow=2, startcol=0, header=False)
        text.to_excel(writer, sheet_name='sheetName', index=False, startrow=3, startcol=0, header=False)
        data1 = df.iloc[i:(unique_values[x]), 4]
        data1 = data1.drop_duplicates()
        data1.to_excel(writer, sheet_name='sheetName', index=False, startrow=3, startcol=1, header=False)
        text2.to_excel(writer, sheet_name='sheetName', index=False, startrow=3, startcol=3, header=False)
        data2 = df.iloc[i:(unique_values[x]), 5]
        data2 = data2.drop_duplicates()
        data2.to_excel(writer, sheet_name='sheetName', index=False, startrow=3, startcol=4, header=False)
        df2 = df.iloc[i:(unique_values[x] - 1), [8, 9, 11, 2, 13, 17]]
        df2.columns = ['Part Number', 'Hardware', 'Qty', 'Serial Number', 'Net Price', 'Net Value']
        df2 = df2.assign(Net_Price = df2.groupby('Serial Number')['Net Price'].transform('sum')).drop('Net Price', axis=1)
        df2 = df2.assign(Net_Value = df2.groupby('Serial Number')['Net Value'].transform('sum')).drop('Net Value', axis=1)
        df2 = df2.drop_duplicates(subset=['Serial Number'])
        print(df2)
        df2.to_excel(writer, sheet_name='sheetName', index=False, startrow=4, startcol=0)
        for column in df2:
            column_length = max(df2[column].astype(str).map(len).max(), len(column))
            col_idx = df2.columns.get_loc(column)
            writer.sheets['sheetName'].set_column(col_idx, col_idx, column_length)

        t = doc.add_table(df2.shape[0]+4, 6)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = True
        t.allow_autofit = True
        t.style = 'Table Grid'
        t.cell(0, 0).text = str(df.iloc[unique_values[x - 1], 16])
        t.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        a = t.cell(0, 0)
        b = t.cell(0, 5)
        a.merge(b)
        t.cell(1, 0).text = str(df.iloc[unique_values[x - 2], 9])
        t.cell(1, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        a = t.cell(1, 0)
        b = t.cell(1, 5)
        a.merge(b)
        t.cell(2, 0).text = "Vigência - De: " + str(df.iloc[unique_values[x - 1], 4]) + " até " + str(df.iloc[unique_values[x - 1], 5])
        t.cell(2, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        a = t.cell(2, 0)
        b = t.cell(2, 5)
        a.merge(b)
        change_table_row(t.rows[0], background_color="#92D050", bold=True)
        change_table_row(t.rows[1], background_color="#92D050", bold=True)
        change_table_row(t.rows[2], background_color="#92D050", bold=True)
        for j in range(df.shape[-1]):
            t.cell(3, 0).text = str("Part Number")
            t.cell(3, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.cell(3, 1).text = str("Hardware")
            t.cell(3, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.cell(3, 2).text = str("Qty")
            t.cell(3, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.cell(3, 3).text = str("Serial Number")
            t.cell(3, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.cell(3, 4).text = str("Net Price")
            t.cell(3, 4).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.cell(3, 5).text = str("Net Value")
            t.cell(3, 5).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for k in range(df2.shape[0]):
                t.cell(k+4,0).text = str(df2.values[k, 0])
                t.cell(k+4, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        for k in range(df2.shape[0]):
                t.cell(k+4,1).text = str(df2.values[k, 1])
                t.cell(k+4, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        for k in range(df2.shape[0]):
                t.cell(k+4,2).text = str(df2.values[k, 2])
                t.cell(k+4, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        for k in range(df2.shape[0]):
                t.cell(k+4,3).text = str(df2.values[k, 3])
                t.cell(k+4, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                if t.cell(k+4,3).text == 'nan':
                     t.cell(k+4,3).text = ' '

        for k in range(df2.shape[0]):
                t.cell(k+4,4).text = str(df2.values[k, 4])
                t.cell(k+4, 4).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        for k in range(df2.shape[0]):
                t.cell(k+4,5).text = str(df2.values[k, 5])
                t.cell(k+4, 5).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        for k in range(df2.shape[0]):
            t.cell(k,3).width = Pt(20)
        
        for k in range(df2.shape[0]):
            for j in range(5):
                if j == 0:
                    t.cell(k,j).width = Mm(100)
                if j == 1:
                    t.cell(k,j).width = Mm(150)
                if j == 2:
                    t.cell(k,j).width = Mm(100)
                elif j == 3:
                    t.cell(k,j).width = Mm(15)
        
        ptable = doc.add_paragraph('\n\n')

        print(unique_values)
        lenght = len(df2)

    if i > 1 and i != unique_values[-1]:
        x += 1
        fl = df.iloc[i:(unique_values[x]), 16]
        fl = fl.drop_duplicates()
        fl.to_excel(writer, sheet_name='sheetName', index=False, startrow=1+lenght+6, startcol=0, header=False)
        sla = df.iloc[i-1:(unique_values[x]), 9]
        sla = sla.drop_duplicates()
        sla.to_excel(writer, sheet_name='sheetName', index=False, startrow=1+lenght+7, startcol=0, header=False)
        text.to_excel(writer, sheet_name='sheetName', index=False, startrow=1+lenght+8, startcol=0, header=False)
        data1 = df.iloc[i:(unique_values[x]), 4]
        data1 = data1.drop_duplicates()
        data1.to_excel(writer, sheet_name='sheetName', index=False, startrow=1+lenght+8, startcol=1, header=False)
        text2.to_excel(writer, sheet_name='sheetName', index=False, startrow=1+lenght+8, startcol=3, header=False)
        data2 = df.iloc[i:(unique_values[x]), 5]
        data2 = data2.drop_duplicates()
        data2.to_excel(writer, sheet_name='sheetName', index=False, startrow=1+lenght+8, startcol=4, header=False)
        y = 1 + lenght + 7
        merge_df.append(y)
        merge_df.append(y+1)
        print(merge_df)
        df2 = df.iloc[i:(unique_values[x] - 1), [8, 9, 11, 2, 13, 17]]
        df2.columns = ['Part Number', 'Hardware', 'Qty', 'Serial Number', 'Net Price', 'Net Value']
        df2 = df2.drop_duplicates(subset=['Serial Number'])
        df2.to_excel(writer, sheet_name='sheetName', index=False, startrow=1+lenght+9, startcol=0)
        for column in df2:
            column_length = max(df2[column].astype(str).map(len).max(), len(column))
            col_idx = df2.columns.get_loc(column)
            writer.sheets['sheetName'].set_column(col_idx, col_idx, column_length)
        lenght = lenght + (len(df2)) + 8

        t = doc.add_table(df2.shape[0]+4, 6)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = True
        t.allow_autofit = True
        t.style = 'Table Grid'
        t.cell(0, 0).text = str(df.iloc[unique_values[x - 1], 16])
        t.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        a = t.cell(0, 0)
        b = t.cell(0, 5)
        a.merge(b)
        t.cell(1, 0).text = str(df.iloc[unique_values[x - 2], 9])
        t.cell(1, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        a = t.cell(1, 0)
        b = t.cell(1, 5)
        a.merge(b)
        t.cell(2, 0).text = "Vigência - De: " + str(df.iloc[unique_values[x - 1], 4]) + " até " + str(df.iloc[unique_values[x - 1], 5])
        t.cell(2, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        a = t.cell(2, 0)
        b = t.cell(2, 5)
        a.merge(b)
        change_table_row(t.rows[0], background_color="#92D050", bold=True)
        change_table_row(t.rows[1], background_color="#92D050", bold=True)
        change_table_row(t.rows[2], background_color="#92D050", bold=True)
        for j in range(df.shape[-1]):
            t.cell(3, 0).text = str("Part Number")
            t.cell(3, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.cell(3, 1).text = str("Hardware")
            t.cell(3, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.cell(3, 2).text = str("Qty")
            t.cell(3, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.cell(3, 3).text = str("Serial Number")
            t.cell(3, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.cell(3, 4).text = str("Net Price")
            t.cell(3, 4).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.cell(3, 5).text = str("Net Value")
            t.cell(3, 5).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for k in range(df2.shape[0]):
                t.cell(k+4,0).text = str(df2.values[k, 0])
                t.cell(k+4, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        for k in range(df2.shape[0]):
                t.cell(k+4,1).text = str(df2.values[k, 1])
                t.cell(k+4, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        for k in range(df2.shape[0]):
                t.cell(k+4,2).text = str(df2.values[k, 2])
                t.cell(k+4, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        for k in range(df2.shape[0]):
                t.cell(k+4,3).text = str(df2.values[k, 3])
                t.cell(k+4, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                if t.cell(k+4,3).text == 'nan':
                     t.cell(k+4,3).text = ' '

        for k in range(df2.shape[0]):
                t.cell(k+4,4).text = str(df2.values[k, 4])
                t.cell(k+4, 4).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        for k in range(df2.shape[0]):
                t.cell(k+4,5).text = str(df2.values[k, 5])
                t.cell(k+4, 5).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        for k in range(df2.shape[0]):
            t.cell(k,3).width = Pt(20)
        
        for k in range(df2.shape[0]):
            for j in range(5):
                if j == 0:
                    t.cell(k,j).width = Mm(100)
                if j == 1:
                    t.cell(k,j).width = Mm(150)
                if j == 2:
                    t.cell(k,j).width = Mm(100)
                elif j == 3:
                    t.cell(k,j).width = Mm(15)
        
        ptable = doc.add_paragraph('\n\n')

        print(lenght)

            
    if i == unique_values[-1]:
        fl = df.iloc[i:, 16]
        fl = fl.drop_duplicates()
        fl.to_excel(writer, sheet_name='sheetName', index=False, startrow=1+lenght+6, startcol=0, header=False)
        sla = df.iloc[i-1:, 9]
        sla = sla.drop_duplicates()
        sla.to_excel(writer, sheet_name='sheetName', index=False, startrow=1+lenght+7, startcol=0, header=False)
        text.to_excel(writer, sheet_name='sheetName', index=False, startrow=1+lenght+8, startcol=0, header=False)
        data1 = df.iloc[i:, 4]
        data1 = data1.drop_duplicates()
        data1.to_excel(writer, sheet_name='sheetName', index=False, startrow=1+lenght+8, startcol=1, header=False)
        text2.to_excel(writer, sheet_name='sheetName', index=False, startrow=1+lenght+8, startcol=3, header=False)
        data2 = df.iloc[i:, 5]
        data2 = data2.drop_duplicates()
        data2.to_excel(writer, sheet_name='sheetName', index=False, startrow=1+lenght+8, startcol=4, header=False)
        y = 1 + lenght + 7
        merge_df.append(y)
        merge_df.append(y+1)
        print(merge_df)
        df2 = df.iloc[i:, [8, 9, 11, 2, 13, 17]]
        df2.columns = ['Part Number', 'Hardware', 'Qty', 'Serial Number', 'Net Price', 'Net Value']
        df2 = df2.drop_duplicates(subset=['Serial Number'])
        df2.to_excel(writer, sheet_name='sheetName', index=False, startrow=1+lenght+9, startcol=0)
        for column in df2:
            column_length = max(df2[column].astype(str).map(len).max(), len(column))
            col_idx = df2.columns.get_loc(column)
            writer.sheets['sheetName'].set_column(col_idx, col_idx, column_length)

        t = doc.add_table(df2.shape[0]+4, 6)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = True
        t.allow_autofit = True
        t.style = 'Table Grid'
        t.cell(0, 0).text = str(df.iloc[unique_values[x - 1], 16])
        t.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        a = t.cell(0, 0)
        b = t.cell(0, 5)
        a.merge(b)
        t.cell(1, 0).text = str(df.iloc[unique_values[x - 2], 9])
        t.cell(1, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        a = t.cell(1, 0)
        b = t.cell(1, 5)
        a.merge(b)
        t.cell(2, 0).text = "Vigência - De: " + str(df.iloc[unique_values[x - 1], 4]) + " até " + str(df.iloc[unique_values[x - 1], 5])
        t.cell(2, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        a = t.cell(2, 0)
        b = t.cell(2, 5)
        a.merge(b)
        change_table_row(t.rows[0], background_color="#92D050", bold=True)
        change_table_row(t.rows[1], background_color="#92D050", bold=True)
        change_table_row(t.rows[2], background_color="#92D050", bold=True)
        for j in range(df.shape[-1]):
            t.cell(3, 0).text = str("Part Number")
            t.cell(3, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.cell(3, 1).text = str("Hardware")
            t.cell(3, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.cell(3, 2).text = str("Qty")
            t.cell(3, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.cell(3, 3).text = str("Serial Number")
            t.cell(3, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.cell(3, 4).text = str("Net Price")
            t.cell(3, 4).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.cell(3, 5).text = str("Net Value")
            t.cell(3, 5).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for k in range(df2.shape[0]):
                t.cell(k+4,0).text = str(df2.values[k, 0])
                t.cell(k+4, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        for k in range(df2.shape[0]):
                t.cell(k+4,1).text = str(df2.values[k, 1])
                t.cell(k+4, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        for k in range(df2.shape[0]):
                t.cell(k+4,2).text = str(df2.values[k, 2])
                t.cell(k+4, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        for k in range(df2.shape[0]):
                t.cell(k+4,3).text = str(df2.values[k, 3])
                t.cell(k+4, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                if t.cell(k+4,3).text == 'nan':
                     t.cell(k+4,3).text = ' '

        for k in range(df2.shape[0]):
                t.cell(k+4,4).text = str(df2.values[k, 4])
                t.cell(k+4, 4).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        for k in range(df2.shape[0]):
                t.cell(k+4,5).text = str(df2.values[k, 5])
                t.cell(k+4, 5).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        for k in range(df2.shape[0]):
            t.cell(k,3).width = Pt(20)
        
        for k in range(df2.shape[0]):
            for j in range(5):
                if j == 0:
                    t.cell(k,j).width = Mm(100)
                if j == 1:
                    t.cell(k,j).width = Mm(150)
                if j == 2:
                    t.cell(k,j).width = Mm(100)
                elif j == 3:
                    t.cell(k,j).width = Mm(15)
        
        ptable = doc.add_paragraph('\n\n')

writer.close()

x = 0
for i in merge_df:
    wb = load_workbook("pandas_to_excel.xlsx")
    ws = wb.active
    start_row = merge_df[x]
    print(start_row)
    end_row = merge_df[x]
    start_column = 1
    end_column = 6
    merge_range = f"{ws.cell(row=start_row, column=start_column).coordinate}:{ws.cell(row=end_row, column=end_column).coordinate}"
    ws.merge_cells(merge_range)
    currentCell = ws.cell(start_row, start_column)
    currentCell.alignment = Alignment(horizontal='center')
    currentCell.fill = PatternFill(start_color='00A98B', end_color='00A98B', fill_type="solid")
    currentCell.border = Border(top=thin, left=thin, right=thin, bottom=thin)
    wb.save("pandas_to_excel.xlsx")
    x += 1

wb_col = load_workbook("pandas_to_excel.xlsx")
sheet = wb_col.active
for col in sheet.columns:
     SetLen = 0
     column = col[0].column_letter
     for cell in col:
             if len(str(cell.value)) > SetLen:
                 SetLen = len(str(cell.value))
     set_col_width = SetLen + 5
     sheet.column_dimensions[column].width = set_col_width       
wb_col.save("pandas_to_excel.xlsx")


wb_style = load_workbook('pandas_to_excel.xlsx')
sheet = wb_style.active
last_row = sheet.max_row
for rows in sheet.iter_rows(min_row=1, max_row=last_row, min_col=0, max_col=6):
   for cell in rows:                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                                         
     if cell.value != None:
        cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)
wb_style.save('pandas_to_excel.xlsx')
doc.save('./test.docx')
print('done!')